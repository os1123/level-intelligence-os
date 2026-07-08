from pathlib import Path
from html import escape

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    KeepTogether,
)
from reportlab.pdfgen import canvas

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/omar/Grounded/Projects/Level Intelligence OS")
OUT = ROOT / "official"
RENDERED = OUT / "rendered"

INK = colors.HexColor("#121820")
MUTED = colors.HexColor("#5b6775")
LINE = colors.HexColor("#d8dde3")
DARK = colors.HexColor("#17202b")
ACCENT = colors.HexColor("#0e806c")
YELLOW = colors.HexColor("#c99b2e")
PAPER = colors.HexColor("#f7f7f4")


def ensure_dirs():
    OUT.mkdir(exist_ok=True)
    RENDERED.mkdir(exist_ok=True)


def styles():
    s = getSampleStyleSheet()
    s.add(ParagraphStyle(
        name="TitleXL",
        parent=s["Title"],
        fontName="Helvetica",
        fontSize=42,
        leading=46,
        textColor=INK,
        spaceAfter=12,
    ))
    s.add(ParagraphStyle(
        name="DeckTitle",
        parent=s["Title"],
        fontName="Helvetica",
        fontSize=34,
        leading=38,
        textColor=INK,
        spaceAfter=10,
    ))
    s.add(ParagraphStyle(
        name="H1x",
        parent=s["Heading1"],
        fontName="Helvetica",
        fontSize=24,
        leading=29,
        textColor=INK,
        spaceBefore=10,
        spaceAfter=8,
    ))
    s.add(ParagraphStyle(
        name="H2x",
        parent=s["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=INK,
        spaceBefore=8,
        spaceAfter=5,
    ))
    s.add(ParagraphStyle(
        name="Lead",
        parent=s["BodyText"],
        fontName="Helvetica",
        fontSize=12.5,
        leading=17,
        textColor=MUTED,
        spaceAfter=10,
    ))
    s.add(ParagraphStyle(
        name="BodyX",
        parent=s["BodyText"],
        fontName="Helvetica",
        fontSize=9.4,
        leading=12.2,
        textColor=MUTED,
        spaceAfter=6,
    ))
    s.add(ParagraphStyle(
        name="Small",
        parent=s["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=MUTED,
    ))
    s.add(ParagraphStyle(
        name="Eyebrow",
        parent=s["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=10,
        textColor=ACCENT,
        spaceAfter=8,
    ))
    s.add(ParagraphStyle(
        name="Quote",
        parent=s["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=17,
        textColor=INK,
        leftIndent=10,
        borderColor=ACCENT,
        borderWidth=0,
        borderPadding=8,
        spaceBefore=6,
        spaceAfter=8,
    ))
    return s


S = styles()


def p(text, style="BodyX"):
    return Paragraph(text, S[style])


def table(rows, widths=None, header=True, font_size=8.4):
    converted = []
    for row in rows:
        converted.append([Paragraph(str(cell), S["Small"]) for cell in row])
    t = Table(converted, colWidths=widths, repeatRows=1 if header else 0)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.55, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]
    if header:
        commands.extend([
            ("BACKGROUND", (0, 0), (-1, 0), DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ])
    t.setStyle(TableStyle(commands))
    return t


def bullets(items):
    return [p(f"- {item}", "BodyX") for item in items]


def on_page(c, doc):
    c.saveState()
    c.setFillColor(INK)
    c.setFont("Helvetica", 8.5)
    c.drawString(doc.leftMargin, LETTER[1] - 0.38 * inch, "Level Intelligence OS")
    c.setFillColor(MUTED)
    c.drawRightString(LETTER[0] - doc.rightMargin, LETTER[1] - 0.38 * inch, f"Page {doc.page}")
    c.setStrokeColor(LINE)
    c.line(doc.leftMargin, LETTER[1] - 0.5 * inch, LETTER[0] - doc.rightMargin, LETTER[1] - 0.5 * inch)
    c.restoreState()


def build_dossier_pdf():
    doc = SimpleDocTemplate(
        str(OUT / "Level Intelligence OS - Official Dossier.pdf"),
        pagesize=LETTER,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.55 * inch,
    )
    story = [
        p("OFFICIAL DOSSIER", "Eyebrow"),
        p("Level Intelligence OS", "TitleXL"),
        p("A private construction intelligence operating system for finding, qualifying, routing, and converting opportunities through Level's own operating judgment.", "Lead"),
        Spacer(1, 0.06 * inch),
        p("Executive Summary", "H1x"),
        p("Level Construction already has the hard part: trusted relationships, commercial construction judgment, development instincts, hospitality experience, capital discipline, and a practical sense for which opportunities are worth time. The opportunity is to turn that judgment into a private operating system that Level owns.", "Lead"),
        p("This is not a generic dashboard, chatbot, CRM migration, or forced Procore replacement. It is a working operating layer around Level's business development, preconstruction, project pursuit, development, hospitality, and capital workflows.", "BodyX"),
        p("<b>This is not built to compete with Mercator or clone any vendor. It is built from the ground up around Level's own relationships, standards, RFQs, approvals, markets, capital context, and operating judgment.</b>", "Quote"),
        p("The near-term outcome is cleaner visibility and faster decisions. The long-term outcome is a proprietary operating memory that improves every time Level reviews a signal, pursues an opportunity, wins a job, passes on a poor-fit project, or learns from execution.", "BodyX"),
        PageBreak(),
        p("The Business Problem", "H1x"),
        p("Level's advantage is operational, but much of that advantage still lives in places that software cannot easily use: emails, RFQ packages, attachments, bid notes, permit comments, spreadsheets, project folders, meeting notes, calls, relationship memory, and executive judgment.", "Lead"),
        table([
            ["Pain", "What It Creates"],
            ["Scattered signals", "Early project, permit, planning, ownership, franchise, broker, and RFQ signals arrive in different places."],
            ["Trapped judgment", "Fit, risk, timing, margin, and relationship context live in individual memory and informal notes."],
            ["RFQ friction", "Deadlines, missing requirements, scope questions, and go/no-go reasoning are buried in language-heavy packages."],
            ["Status burden", "Leadership spends time assembling context instead of making decisions."],
        ], [1.6 * inch, 4.85 * inch]),
        Spacer(1, 0.16 * inch),
        p("Outcomes, Not Just Code", "H1x"),
        p("We give you outcomes, not just code. Code is the mechanism. The product is operating lift: cleaner pipeline visibility, faster triage, better RFQ discipline, fewer missed follow-ups, a weekly executive decision rhythm, and reusable enterprise memory.", "Lead"),
        table([
            ["Outcome", "What Changes"],
            ["Pipeline visibility", "Active pursuits have owner, stage, score, source, next action, and age."],
            ["RFQ discipline", "Deadlines, requirements, missing information, and go/no-go reasoning are extracted and routed."],
            ["Decision speed", "Approvals and exceptions surface before they become bottlenecks."],
            ["Executive rhythm", "Leadership gets a weekly brief from live system data."],
            ["Enterprise memory", "Signals, decisions, outcomes, and overrides improve the next cycle."],
        ], [1.55 * inch, 4.9 * inch]),
        PageBreak(),
        p("What The System Does", "H1x"),
        table([
            ["System Area", "Client-Facing Outcome"],
            ["Signal inbox", "Early project, permit, planning, ownership, franchise, broker, and RFQ signals captured in one place."],
            ["Opportunity graph", "Accounts, contacts, sites, jurisdictions, RFQs, bids, permits, and decisions connected."],
            ["Qualification engine", "Pursue, watch, or pass recommendations based on Level-specific fit criteria."],
            ["Relationship memory", "Searchable account/contact context, relationship owners, last touch, and next follow-up."],
            ["RFQ command center", "Scope summaries, deadlines, missing requirements, go/no-go memos, and bid tasks."],
            ["Approval queue", "Consequential actions routed to humans before external movement."],
            ["Executive brief", "Weekly view of what changed, what is stuck, what needs action, and what leadership must decide."],
        ], [1.55 * inch, 4.9 * inch]),
        Spacer(1, 0.16 * inch),
        p("Human Control And Trust", "H1x"),
        table([
            ["Authority Level", "System Can Do", "Human Control"],
            ["Autonomous", "Classify signals, summarize RFQs, dedupe contacts, draft briefs.", "Logged for review."],
            ["Suggest", "Recommend owner, next action, pursue/watch/pass, follow-up draft.", "Human can accept or edit."],
            ["Approval Required", "External outreach, bid submission, vendor award, budget change, investor communication.", "Explicit approval required."],
            ["Manual", "Legal, contract, financing, final strategic decisions.", "Human executes."],
        ], [1.35 * inch, 3.05 * inch, 2.05 * inch]),
        PageBreak(),
        p("Procore-First Bridge", "H1x"),
        p("The recommended first move is to wrap around Procore, not replace it. Procore can remain the execution source of truth while Level Intelligence OS becomes the intelligence, pursuit, RFQ, relationship, approval, and executive layer around it.", "Lead"),
        table([
            ["Now", "Later"],
            ["Read and reference Procore data where it improves visibility, without disrupting active project teams.", "Use observed value to decide which workflows should remain in Procore and which can move into Level-owned software."],
            ["Structure Level Intelligence OS around current tools, current approvals, and current operating reality.", "Reduce monolith dependence module by module after the internal layer proves better fit, lower friction, and clear ROI."],
        ], [3.2 * inch, 3.25 * inch]),
        Spacer(1, 0.16 * inch),
        p("Proven Internal Patterns", "H1x"),
        p("The build can reuse known construction-platform primitives: connected domain models, workflow state machines, approval queues, document extraction, specialist agent handoffs, audit logs, role-specific dashboards, and feedback loops from outcomes.", "BodyX"),
        p("This lowers build risk while keeping the client-facing system fully Level-specific.", "Quote"),
        PageBreak(),
        p("90-Day Pilot", "H1x"),
        p("Objective: prove that Level Intelligence OS improves opportunity visibility, qualification discipline, RFQ coordination, executive reporting, and institutional memory.", "Lead"),
        table([
            ["Period", "Focus", "Deliverable"],
            ["Weeks 1-2", "Foundation", "Data model, roles, active opportunity/account/contact views, executive dashboard skeleton."],
            ["Weeks 3-5", "Intelligence", "Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest."],
            ["Weeks 6-8", "Workflow", "RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking."],
            ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, outcome capture, adoption review, phase-two roadmap."],
        ], [1.2 * inch, 1.6 * inch, 3.65 * inch]),
        Spacer(1, 0.16 * inch),
        p("Pilot Success Metrics", "H2x"),
        *bullets([
            "100 percent of active pursuits visible with owner, stage, source, score, next action, and age.",
            "At least 10 useful signals captured, reviewed, and routed.",
            "RFQs tracked with deadlines, requirements, and go/no-go status.",
            "Weekly executive brief generated from live system data.",
            "Leadership can point to faster decisions, fewer blind spots, or clearer accountability.",
            "Phase-two roadmap is based on observed value, not assumptions.",
        ]),
        PageBreak(),
        p("Commercial And Ownership Principles", "H1x"),
        table([
            ["Area", "Recommended Position"],
            ["Engagement", "90-day pilot for a private, Level-specific construction intelligence operating system."],
            ["Pilot deliverables", "Data model, opportunity/contact database, signal inbox, scoring framework, RFQ workflow, approval queue, weekly executive brief, success report, phase-two roadmap."],
            ["Ownership", "Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. Code ownership/licensing to be defined in the final agreement."],
            ["Confidentiality", "Level data is confidential. No external reuse without written permission. Role-based access and audit logging required."],
            ["Commercial structure", "Fixed-fee pilot, build fee plus monthly operating retainer, strategy/design sprint followed by implementation, or long-term internal software partnership."],
        ], [1.55 * inch, 4.9 * inch]),
        Spacer(1, 0.18 * inch),
        p("Closing Position", "H1x"),
        p("The first version gives Level cleaner visibility and faster decisions. The strategic version gives Level a proprietary operating memory that improves how the company finds, qualifies, wins, and executes work.", "Lead"),
        p("<b>Level should not rent generic intelligence when it can build the operating memory that makes its own team sharper.</b>", "Quote"),
    ]
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def build_comprehensive_packet_pdf():
    doc = SimpleDocTemplate(
        str(OUT / "Level Intelligence OS - Comprehensive Client Packet.pdf"),
        pagesize=LETTER,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.55 * inch,
    )
    story = [
        p("COMPREHENSIVE CLIENT PACKET", "Eyebrow"),
        p("Level Intelligence OS", "TitleXL"),
        p("A private construction intelligence operating system built from the ground up around Level's own relationships, standards, RFQs, approvals, markets, capital context, and operating judgment.", "Lead"),
        Spacer(1, 0.18 * inch),
        p("<b>Commercial promise:</b> We give you outcomes, not just code.", "Quote"),
        p("Code is the mechanism. The product is operating lift: cleaner opportunity visibility, faster RFQ triage, stronger go/no-go discipline, fewer missed follow-ups, a weekly executive decision rhythm, and Level-owned operating memory that compounds over time.", "Lead"),
        Spacer(1, 0.24 * inch),
        table([
            ["Packet Section", "Purpose"],
            ["1. Executive thesis", "Why Level should own a private operating layer."],
            ["2. Pain and opportunity", "Where scattered judgment creates drag."],
            ["3. Built for Level", "Why this is not a vendor clone or competitor pitch."],
            ["4. Proven infrastructure", "Reusable construction-platform primitives that lower build risk."],
            ["5. Procore bridge", "How to wrap around the current source of truth before replacing anything."],
            ["6. Market shift", "Why software is moving toward forward-deployed, company-specific systems."],
            ["7. System design", "How find, qualify, route, convert, and learn becomes software."],
            ["8. Pilot and terms", "How to prove value in 90 days and structure the engagement."],
        ], [1.65 * inch, 4.8 * inch]),
        PageBreak(),

        p("1. Executive Thesis", "H1x"),
        p("Level Construction already has the hard part: trusted relationships, construction judgment, development instincts, hospitality experience, capital discipline, vendor memory, and practical pattern recognition about which opportunities are worth time.", "Lead"),
        p("The opportunity is to convert that judgment into a private operating system that Level owns. Level Intelligence OS is not a generic AI layer, chatbot, dashboard, CRM migration, or forced Procore replacement. It is an internal operating layer around business development, preconstruction, project pursuit, development, hospitality, capital workflows, and executive management.", "BodyX"),
        p("The first version creates visibility and decision speed. The strategic version becomes Level's proprietary construction intelligence layer: a system that learns from Level's signals, relationships, decisions, wins, losses, and operating judgment.", "BodyX"),
        p("<b>Core proposition:</b> Level should not rent generic intelligence when it can build the operating memory that makes its own team sharper.", "Quote"),
        table([
            ["Near-Term Outcome", "Long-Term Strategic Asset"],
            ["Cleaner active pipeline, faster RFQ review, better go/no-go discipline, and weekly executive visibility.", "A Level-owned operating memory that improves pursuit, relationship leverage, execution handoff, and future software adaptability."],
        ], [3.05 * inch, 3.4 * inch]),
        PageBreak(),

        p("2. The Pain Before The Platform", "H1x"),
        p("The highest-value information in construction pursuit work rarely lives in one system. It lives in RFQs, emails, calls, spreadsheets, permit threads, broker notes, project folders, meeting notes, personal memory, and executive judgment.", "Lead"),
        table([
            ["Pain Point", "What It Looks Like", "Why It Matters"],
            ["Scattered signals", "Permits, planning activity, ownership changes, broker notes, franchise news, RFQs, and relationship tips arrive in different places.", "Promising opportunities can be seen late, duplicated, or missed."],
            ["Trapped judgment", "Fit, risk, timing, margin, and relationship context live in individual memory, inboxes, and informal notes.", "Level's advantage is real but hard to search, compare, hand off, or improve."],
            ["RFQ friction", "Deadlines, missing documents, addenda, site risk, scope ambiguity, and submission requirements are buried in language-heavy packages.", "Teams spend time translating documents into tasks and may surface risk late."],
            ["Status burden", "Leadership needs to ask what changed, what is stuck, who owns it, and what needs approval.", "Executive time is spent assembling context instead of exercising judgment."],
        ], [1.3 * inch, 2.7 * inch, 2.45 * inch]),
        Spacer(1, 0.16 * inch),
        p("The result is not a lack of effort. It is a lack of one operating layer that can answer Level's recurring leadership questions quickly.", "BodyX"),
        *bullets([
            "Which opportunities are active right now?",
            "Which are worth pursuing, watching, or passing on?",
            "Which RFQs are due soon, incomplete, or under-reviewed?",
            "Which relationships need follow-up?",
            "Which decisions are waiting on leadership?",
            "Which assumptions keep repeating across wins and losses?",
        ]),
        PageBreak(),

        p("3. Built For Level, Not For The Category", "H1x"),
        p("Mercator and similar tools are useful category references because they show that earlier construction intelligence has value. But Level Intelligence OS is not a Mercator replacement pitch and not a vendor comparison exercise.", "Lead"),
        p("The strategic point is that Level should own software built from the ground up around Level's own relationships, standards, RFQ process, approval rules, development instincts, hospitality pipeline, capital context, and executive operating rhythm.", "BodyX"),
        table([
            ["Generic Vendor Logic", "Level Intelligence OS Logic"],
            ["Expose market or project signals.", "Expose signals and translate them into Level-specific opportunities, owners, scores, and next actions."],
            ["Use vendor-defined fields and workflows.", "Represent Level's own domain model: accounts, contacts, sites, RFQs, permits, bids, approvals, relationships, and decisions."],
            ["Optimize for broad category adoption.", "Optimize for Level's operating judgment, pursuit discipline, leadership rhythm, and long-term ownership."],
            ["Data remains mainly a tool output.", "Data becomes Level-owned operating memory that compounds over time."],
        ], [2.75 * inch, 3.7 * inch]),
        Spacer(1, 0.18 * inch),
        p("<b>Positioning sentence:</b> Level Intelligence OS is not built to compete with anyone else's product. It is built to make Level's own operating model software-native.", "Quote"),
        PageBreak(),

        p("4. Proven Infrastructure Patterns We Bring", "H1x"),
        p("We are not starting from an empty page. The build can reuse internal construction-platform patterns already worked through in prior platform development: data models, workflow queues, approval gates, agent handoffs, document extraction, audit trails, dashboard surfaces, and feedback loops.", "Lead"),
        table([
            ["Pattern", "How It Applies To Level"],
            ["Construction domain model", "Accounts, contacts, sites, projects, permits, RFQs, bids, budget items, approvals, outcomes, and operating notes can be represented as connected records."],
            ["Workflow state machines", "Opportunities, RFQs, permits, approvals, and later payment or vendor workflows move through explicit statuses instead of living as scattered messages."],
            ["Specialist agent structure", "Orchestrated agents handle bounded jobs such as signal intake, qualification, RFQ extraction, permit monitoring, executive briefing, and finance context."],
            ["Human-in-the-loop controls", "Authority levels, approval queues, audit logs, escalation rules, and explicit human review keep consequential actions under Level control."],
            ["Document intelligence", "RFQs, permits, invoices, emails, addenda, and project narratives can be transformed into deadlines, risks, tasks, missing information, and decision memos."],
            ["Operating dashboards", "Executives, pursuit owners, RFQ reviewers, and future project teams get role-specific views instead of one static dashboard."],
            ["Feedback loops", "Wins, losses, overrides, delays, vendor outcomes, and cycle times become training signals for better future recommendations."],
        ], [1.75 * inch, 4.7 * inch]),
        Spacer(1, 0.16 * inch),
        p("The point is not to expose any prior build. The point is that the primitives are already known, which gives Level a faster, lower-risk path to a private system.", "Quote"),
        PageBreak(),

        p("5. Procore Bridge Strategy", "H1x"),
        p("Level Intelligence OS should wrap around Procore first, not replace it. Procore is a strong source of truth for active project execution. The first job is to structure the new platform around that reality while adding the intelligence, pursuit, RFQ, relationship, approval, and executive layer that Procore is not designed to be.", "Lead"),
        table([
            ["Phase", "Position", "Outcome"],
            ["1. Read and reference", "Treat Procore as the current project source of truth. Pull or reference only the data needed for visibility and decisions.", "No disruption to field or project teams."],
            ["2. Workflow overlay", "Use Level Intelligence OS for signals, RFQs, go/no-go memos, relationship memory, approvals, and executive briefs.", "Level gains intelligence and operating rhythm without forcing a platform migration."],
            ["3. Evidence-based rationalization", "Measure which workflows are better handled inside Level's own system and which should stay in Procore.", "Software decisions become based on observed value, not vendor ideology."],
            ["4. Selective replacement", "Move modules away from monolithic software only when Level's owned layer is clearly better, safer, and cheaper to operate.", "Dependency reduces over time without risky rip-and-replace."],
        ], [1.35 * inch, 2.65 * inch, 2.45 * inch]),
        Spacer(1, 0.16 * inch),
        p("<b>Operating principle:</b> earn the right to replace. The long-term goal is less reliance on monolithic software, but the first move is a Procore-first intelligence layer that creates measurable value now.", "Quote"),
        PageBreak(),

        p("6. Market Shift: Forward-Deployed Software", "H1x"),
        p("The software market is moving from fixed SaaS screens toward company-specific AI systems deployed inside real workflows. OpenAI's 2026 Deployment Company and forward-deployed engineering model are a clear market signal: enterprise value increasingly comes from fitting intelligence into real operating constraints, not simply licensing a generic tool.", "Lead"),
        table([
            ["Shift", "Old Pattern", "New Pattern"],
            ["Static SaaS to internal systems", "The company adapts to vendor-defined screens and workflows.", "Software adapts to the company's language, judgment, data, approvals, and operating rhythm."],
            ["Generic automation to agentic work", "Rules handle narrow tasks but struggle with messy language-heavy work.", "Agents read, synthesize, route, draft, and learn from RFQs, emails, permits, notes, and decisions."],
            ["Vendor lock-in to owned memory", "Operating history lives inside external tools and exports.", "The company's own decisions, outcomes, and workflows become a proprietary asset."],
            ["Implementation to forward deployment", "Software is installed and handed off.", "Experts map work, ship inside real constraints, tune with operators, and turn adoption into measurable outcomes."],
            ["Manual change requests to language-native change", "New workflows require slow product cycles or vendor roadmaps.", "Teams can describe additions, omissions, and changes in natural language, then convert them into owned software modules."],
        ], [1.4 * inch, 2.4 * inch, 2.65 * inch]),
        Spacer(1, 0.16 * inch),
        p("For Level, this makes the proposal timely: build the internal layer now, compound proprietary memory, and let that memory guide which vendor systems remain essential versus replaceable.", "Quote"),
        PageBreak(),

        p("7. What The System Does", "H1x"),
        p("Level Intelligence OS converts messy pursuit work into an operating loop: find, qualify, route, convert, and learn.", "Lead"),
        table([
            ["Capability", "What It Does", "Client-Facing Outcome"],
            ["Find", "Detects public and private project signals before they become obvious.", "Level sees more opportunities earlier."],
            ["Qualify", "Scores opportunities using Level's criteria, history, relationships, and capacity.", "Leadership spends less time on poor-fit work."],
            ["Route", "Assigns owners, next actions, approvals, and follow-ups.", "Good opportunities do not stall in inboxes."],
            ["Convert", "Supports RFQs, go/no-go memos, pursuit briefs, and executive reporting.", "Level turns intelligence into action."],
            ["Learn", "Captures decisions, overrides, wins, losses, delays, and outcomes.", "Every cycle improves the next score, brief, and recommendation."],
        ], [1.1 * inch, 2.75 * inch, 2.6 * inch]),
        Spacer(1, 0.18 * inch),
        p("The goal is not to remove judgment. The goal is to put better context in front of the people making the judgment, then preserve what happened so the system and organization get sharper.", "BodyX"),
        PageBreak(),

        p("8. Product Modules", "H1x"),
        table([
            ["Module", "What It Includes", "Why It Matters"],
            ["Signal inbox", "Permits, planning agendas, title/ownership changes, rezoning, franchise expansion, broker notes, inbound RFQs, CRM/email activity.", "Early signals become reviewable records instead of scattered noise."],
            ["Opportunity graph", "Accounts, contacts, sites, jurisdictions, RFQs, bids, permits, decisions, owners, next actions.", "Level sees relationships between people, projects, sites, and pursuits."],
            ["Qualification engine", "Pursue/watch/pass, confidence, reasoning, missing info, risks, next action, owner, comparable history.", "Level's own standards become repeatable without replacing human judgment."],
            ["Relationship memory", "Relationship owner, last touch, prior work, warm intros, next follow-up, leadership brief.", "Relationship context becomes searchable and actionable."],
            ["RFQ command center", "Scope summary, deadlines, addenda, required attachments, missing documents, risks, go/no-go memo, bid tasks.", "RFQs become decision-ready work."],
            ["Approval queue", "Recommendations, evidence, confidence, approver, decision, outcome.", "Agents can prepare work while humans control consequential actions."],
            ["Executive brief", "Top signals, active pursuits, stalled items, RFQs due, approvals needed, relationship follow-ups, lessons learned.", "Leadership gets a weekly decision rhythm from live system data."],
        ], [1.25 * inch, 3.05 * inch, 2.15 * inch]),
        PageBreak(),

        p("9. Agentic Operating Model", "H1x"),
        p("The system should be composed of specialist workflows under one governance model. This avoids the failure mode of one vague assistant that is expected to know everything.", "Lead"),
        table([
            ["Workflow Agent", "Role"],
            ["Signal Scout", "Finds, dedupes, summarizes, and routes project signals."],
            ["Qualification", "Scores opportunities against Level-specific criteria and explains the recommendation."],
            ["Relationship", "Builds account/contact briefs, tracks owners, suggests follow-ups, and supports warm intros."],
            ["RFQ/Bid", "Extracts requirements, deadlines, risks, missing documents, addenda, and go/no-go memos."],
            ["Permit/Entitlement", "Tracks jurisdiction status, review comments, days in review, and hidden timing risk."],
            ["Executive Briefing", "Turns system activity into a weekly decision brief."],
            ["Finance/Capital", "Connects opportunities to financing, underwriting, investor context, and capital constraints where relevant."],
        ], [1.6 * inch, 4.85 * inch]),
        Spacer(1, 0.16 * inch),
        table([
            ["Authority Level", "System Can Do", "Human Control"],
            ["Autonomous", "Classify signals, summarize RFQs, dedupe contacts, draft internal briefs.", "Logged for review."],
            ["Suggest", "Recommend owner, next action, pursue/watch/pass, follow-up draft.", "Human can accept, edit, or reject."],
            ["Approval Required", "External outreach, bid submission, vendor award, budget change, investor communication.", "Explicit approval required."],
            ["Manual", "Legal, contracts, financing, final strategic decisions.", "Human executes."],
        ], [1.35 * inch, 3.0 * inch, 2.1 * inch]),
        PageBreak(),

        p("10. ROI And Value Model", "H1x"),
        p("The ROI should not be sold only as labor savings. The bigger value is better pursuit discipline, faster decisions, fewer blind spots, stronger relationship leverage, and proprietary operating memory.", "Lead"),
        table([
            ["Value Lever", "How Value Is Created", "How To Measure"],
            ["Opportunity capture", "Signals become structured opportunities with owner, score, source, and next action.", "Signals captured, time to first action, conversion by source, wins by source."],
            ["Pursuit discipline", "Low-fit pursuits are filtered earlier while high-fit work receives faster attention.", "Go/no-go cycle time, win rate by score, avoided pursuit hours, margin by source."],
            ["Admin reduction", "Agents summarize, extract, draft, route, and remind across language-heavy workflows.", "Hours saved per RFQ, weekly report time saved, overdue task reduction."],
            ["Risk reduction", "Permits, missing documents, stalled decisions, and deadline risk surface sooner.", "Missed deadline rate, response time, permits past expected review, unresolved blocker age."],
            ["Enterprise asset", "Relationships, decisions, outcomes, vendors, and market patterns become private memory.", "Record completeness, reuse of history, recommendation accuracy, override rate."],
        ], [1.4 * inch, 3.0 * inch, 2.05 * inch]),
        Spacer(1, 0.16 * inch),
        p("<b>Compounding logic:</b> each signal, decision, win, loss, permit delay, vendor outcome, and relationship touch improves the next recommendation.", "Quote"),
        PageBreak(),

        p("11. 90-Day Pilot", "H1x"),
        p("The pilot should prove usefulness in workflows Level already runs: opportunity review, RFQ triage, relationship follow-up, approval routing, and executive reporting.", "Lead"),
        table([
            ["Period", "Focus", "Client-Visible Deliverable"],
            ["Weeks 1-2", "Foundation", "Data model, roles, active opportunities, account/contact layer, executive dashboard skeleton."],
            ["Weeks 3-5", "Intelligence", "Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest."],
            ["Weeks 6-8", "Workflow", "RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking."],
            ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, outcome tracking, adoption review, phase-two roadmap."],
        ], [1.15 * inch, 1.6 * inch, 3.7 * inch]),
        Spacer(1, 0.16 * inch),
        p("Pilot Success Metrics", "H2x"),
        *bullets([
            "100 percent of active pursuits visible in the system.",
            "Each pursuit has owner, stage, source, score, next action, and age.",
            "At least 10 useful signals captured, reviewed, and routed.",
            "RFQs tracked with deadlines, requirements, and go/no-go status.",
            "Weekly executive brief generated from live system data.",
            "Leadership identifies faster decisions, fewer blind spots, or clearer accountability.",
            "Phase-two roadmap is based on observed value, not assumptions.",
        ]),
        PageBreak(),

        p("12. Phase Two And System Expansion", "H1x"),
        p("Expansion should be based on where Level sees the highest operating value during the pilot, not a preset software wish list.", "Lead"),
        table([
            ["Expansion Area", "Potential Scope"],
            ["Permit and entitlement monitoring", "Jurisdiction status, comments, timing risk, follow-up drafts, days in review."],
            ["Email, calendar, CRM, storage integration", "Source access, relationship context, meeting history, follow-up capture."],
            ["Procore read-only bridge", "Treat Procore as the current project source of truth while Level Intelligence OS adds intelligence, pursuit, and executive workflow above it."],
            ["Vendor and subcontractor memory", "Performance history, issue patterns, pricing, reliability, project outcomes."],
            ["Project handoff", "Transfer pursuit context into execution planning and kickoff."],
            ["Finance and capital workflow", "Underwriting context, investor notes, financing constraints, capital approval support."],
            ["Hospitality pipeline", "Brand/operator tracking, franchise expansion, site qualification, relationship history."],
            ["Software rationalization", "Usage audit to decide what should remain, integrate, downgrade, or eventually be replaced by Level-owned modules."],
        ], [2.0 * inch, 4.45 * inch]),
        Spacer(1, 0.14 * inch),
        p("The recommended position is not 'replace Procore' at the start. The stronger position is to wrap around it, learn where the friction is, and use evidence to decide which workflows can eventually move into Level-owned software.", "BodyX"),
        PageBreak(),

        p("13. Commercial, Ownership, And Trust Principles", "H1x"),
        table([
            ["Area", "Recommended Position"],
            ["Engagement", "90-day pilot to design, build, test, and operationalize Level Intelligence OS v1."],
            ["Pilot deliverables", "Workflow map, data model, opportunity/account/contact database, signal inbox, scoring framework, RFQ workflow, approval queue, weekly executive brief, success report, phase-two roadmap."],
            ["Ownership", "Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. Code ownership/licensing to be defined in the final agreement."],
            ["Model infrastructure", "AI model providers remain replaceable infrastructure, not the strategic asset."],
            ["Confidentiality", "Level data is confidential. No external reuse without written permission. Role-based access and audit logs are required."],
            ["Human approval", "External outreach, bid submission, vendor award, budget change, investor communication, legal, financing, and final strategic decisions require human control."],
            ["Commercial options", "Fixed-fee pilot, build fee plus monthly operating retainer, strategy/design sprint followed by implementation, or long-term internal software partnership."],
        ], [1.55 * inch, 4.9 * inch]),
        PageBreak(),

        p("14. Discovery Agenda", "H1x"),
        p("The first discovery session should focus on actual workflow, not abstract AI strategy.", "Lead"),
        table([
            ["Discovery Topic", "Questions To Answer"],
            ["Opportunity sources", "Where do new opportunities originate today? Which sources matter most?"],
            ["Qualification criteria", "What makes a Level opportunity attractive or unattractive? Who decides?"],
            ["RFQ workflow", "Where do RFQs live? How are deadlines, addenda, tasks, and go/no-go decisions handled?"],
            ["Relationship memory", "Which accounts, contacts, owners, brands, developers, and vendors matter most?"],
            ["Executive reporting", "What does leadership need to see every week? What is currently hard to assemble?"],
            ["Systems and data", "Which tools are in use today: Procore, CRM, email, calendar, accounting, storage? What should be read, synced, imported, or left alone?"],
            ["Approval rules", "Which actions require explicit approval? Who approves them?"],
            ["Pilot success", "What would make the pilot obviously worth continuing?"],
        ], [1.65 * inch, 4.8 * inch]),
        Spacer(1, 0.18 * inch),
        p("Recommended next step: run a 60-minute workflow discovery session, collect sample RFQs and current active pursuits, then finalize the pilot scope and commercial structure.", "Quote"),
        PageBreak(),

        p("15. Closing Position", "H1x"),
        p("Level Intelligence OS should be presented as a serious operating asset, not an AI experiment. The first version gives Level one place to see opportunities, understand fit, track RFQs, preserve relationship memory, and route decisions.", "Lead"),
        p("The longer-term version becomes Level's private construction intelligence layer: a system that learns from Level's signals, relationships, decisions, wins, losses, and operating judgment.", "BodyX"),
        p("<b>Final message:</b> Level can build internal software that grows with the company, adapts through language, and compounds from its own operating history.", "Quote"),
        table([
            ["Immediate Ask", "Approve a focused 90-day pilot."],
            ["What Level Gets", "A working operating rhythm, not a codebase sitting on a shelf."],
            ["Why It Matters", "The software becomes more valuable because it learns Level."],
        ], [1.65 * inch, 4.8 * inch]),
    ]
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def build_one_page_pdf():
    path = OUT / "Level Intelligence OS - Official One Pager.pdf"
    c = canvas.Canvas(str(path), pagesize=landscape(LETTER))
    pw, ph = landscape(LETTER)
    m = 0.35 * inch
    cw = pw - 2 * m
    y = ph - m

    def draw_box(x, top, w, h, fill=None, stroke=LINE):
        c.setStrokeColor(stroke)
        c.setFillColor(fill or colors.white)
        c.rect(x, top - h, w, h, stroke=True, fill=True)

    def wrap(text, max_chars):
        words, lines, line = text.split(), [], ""
        for word in words:
            test = f"{line} {word}".strip()
            if len(test) <= max_chars:
                line = test
            else:
                lines.append(line)
                line = word
        if line:
            lines.append(line)
        return lines

    def text(text, x, top, w, size=8.5, leading=10.2, color=MUTED, bold=False, max_lines=8):
        c.setFillColor(color)
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        chars = max(18, int(w / (size * 0.45)))
        yy = top
        for line in wrap(text, chars)[:max_lines]:
            c.drawString(x, yy, line)
            yy -= leading
        return yy

    c.setFillColor(ACCENT)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(m, y, "OFFICIAL ONE-PAGE DOSSIER")
    y -= 0.36 * inch
    c.setFillColor(INK)
    c.setFont("Helvetica", 35)
    c.drawString(m, y, "Level Intelligence OS")
    c.setStrokeColor(INK)
    c.line(m, y - 14, pw - m, y - 14)
    c.rect(pw - m - 45, y - 2, 38, 32, stroke=True, fill=False)
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(pw - m - 26, y + 8, "LX")
    y -= 0.34 * inch

    draw_box(m, y, cw, 64)
    c.setFillColor(INK)
    c.setFont("Helvetica", 15)
    c.drawString(m + 10, y - 18, "The pain points")
    cols = [
        ("Scattered signals", "Project, permit, planning, broker, franchise, relationship, and RFQ signals arrive in different places."),
        ("Trapped judgment", "Fit, risk, timing, margin, and relationship context live in memory, inboxes, and informal notes."),
        ("Status burden", "Leadership needs what changed, what is stuck, what needs approval, and who owns the next move."),
    ]
    col_w = (cw - 36) / 3
    for i, (h, b) in enumerate(cols):
        x = m + 10 + i * (col_w + 8)
        c.setStrokeColor(INK)
        c.line(x, y - 30, x + col_w, y - 30)
        text(h, x, y - 43, col_w, size=8.2, color=INK, bold=True, max_lines=1)
        text(b, x, y - 54, col_w, size=7.1, leading=8.2, max_lines=2)
    y -= 73

    left_w = cw * 0.49
    right_w = cw - left_w - 10
    draw_box(m, y, left_w, 86, fill=DARK, stroke=DARK)
    text("Outcomes, not just code", m + 10, y - 19, left_w - 20, size=15, color=colors.white, bold=False, max_lines=1)
    text("Code is the mechanism. The product is operating lift: cleaner pipeline visibility, faster RFQ triage, fewer missed follow-ups, weekly executive decision rhythm, and reusable operating memory.", m + 10, y - 40, left_w - 20, size=8, leading=9.3, color=colors.HexColor("#dce3e8"), max_lines=4)
    draw_box(m + left_w + 10, y, right_w, 86)
    text("Built for Level", m + left_w + 20, y - 19, right_w - 20, size=15, color=INK, max_lines=1)
    text("This is not a vendor clone or competitor pitch. It wraps around Procore first and is built around Level's relationships, standards, RFQs, approvals, markets, capital context, and operating history.", m + left_w + 20, y - 40, right_w - 30, size=8, leading=9.3, color=MUTED, max_lines=5)
    y -= 96

    draw_box(m, y, cw, 112)
    text("What Level gets", m + 10, y - 19, cw - 20, size=15, color=INK, max_lines=1)
    gets = [
        ("Signal inbox", "Early project, permit, planning, ownership, franchise, broker, and RFQ signals."),
        ("Qualification engine", "Pursue, watch, or pass recommendations based on Level-specific fit criteria."),
        ("Relationship memory", "Account/contact context, relationship owner, last touch, and next follow-up."),
        ("RFQ command center", "Scope summaries, deadlines, missing requirements, go/no-go memos, and bid tasks."),
        ("Approval queue", "Consequential actions routed to humans before external movement."),
        ("Executive brief", "Weekly view of what changed, what is stuck, and what leadership must decide."),
    ]
    cell_w = (cw - 20 - 5 * 8) / 6
    for i, (h, b) in enumerate(gets):
        x = m + 10 + i * (cell_w + 8)
        c.setStrokeColor(INK)
        c.line(x, y - 34, x + cell_w, y - 34)
        text(h, x, y - 47, cell_w, size=7.7, color=INK, bold=True, max_lines=2)
        text(b, x, y - 61, cell_w, size=6.6, leading=7.4, max_lines=5)
    y -= 122

    half = (cw - 10) / 2
    draw_box(m, y, half, 110)
    text("90-day pilot", m + 10, y - 19, half - 20, size=15, color=INK, max_lines=1)
    pilot = [
        ("Weeks 1-2", "Foundation: data model, roles, active pursuits, contacts, executive view."),
        ("Weeks 3-5", "Intelligence: signal inbox, summaries, scoring, relationship briefs."),
        ("Weeks 6-8", "Workflow: RFQ command center, go/no-go memo, approval queue."),
        ("Weeks 9-12", "Rhythm: weekly brief, outcome capture, adoption review, roadmap."),
    ]
    yy = y - 38
    for h, b in pilot:
        text(h, m + 10, yy, 70, size=7.2, color=INK, bold=True, max_lines=1)
        yy = text(b, m + 92, yy, half - 110, size=7.1, leading=8, max_lines=2) - 2

    draw_box(m + half + 10, y, half, 110, fill=DARK, stroke=DARK)
    text("Term sheet skeleton", m + half + 20, y - 19, half - 20, size=15, color=colors.white, max_lines=1)
    terms = [
        ("Engagement", "90-day pilot for Level Intelligence OS v1."),
        ("Deliverables", "Data model, Procore-first bridge, signal inbox, scoring, RFQ workflow, approval queue, weekly brief, success report."),
        ("Ownership", "Level owns data, operating memory, domain model, workflow definitions, and exportable records."),
        ("Commercial", "Fixed-fee pilot, build plus retainer, strategy sprint, or long-term internal software partnership."),
    ]
    yy = y - 39
    for h, b in terms:
        text(h, m + half + 20, yy, 70, size=7.1, color=colors.white, bold=True, max_lines=1)
        yy = text(b, m + half + 98, yy, half - 116, size=6.9, leading=7.7, color=colors.HexColor("#dce3e8"), max_lines=2) - 2

    c.setFillColor(MUTED)
    c.setFont("Helvetica", 7.5)
    c.drawString(m, 14, "Level Intelligence OS - official one-page dossier")
    c.drawRightString(pw - m, 14, "Draft for discussion")
    c.save()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def docx_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor(58, 69, 81)
    for name, size in [("Heading 1", 21), ("Heading 2", 14), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(17, 22, 29)


def add_docx_table(doc, rows):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, val in enumerate(rows[0]):
        t.rows[0].cells[i].text = str(val)
        set_cell_shading(t.rows[0].cells[i], "17202B")
        set_cell_text_color(t.rows[0].cells[i], "FFFFFF")
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


def add_title(doc, title, subtitle):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(title + "\n")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(17, 22, 29)
    sub = para.add_run(subtitle)
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(13, 128, 108)


def build_docx_files():
    one = Document()
    docx_styles(one)
    sec = one.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)
    add_title(one, "Level Intelligence OS", "Official one-page dossier and term sheet")
    one.add_paragraph("A private construction intelligence operating system for finding, qualifying, routing, and converting opportunities through Level's own operating judgment.")
    add_docx_table(one, [["Pain Points"], ["Signals are scattered across permits, planning, broker notes, franchise expansion, RFQs, and relationships. Judgment is trapped in people, inboxes, and project memory. Leadership needs exceptions, decisions, risks, and next actions in one place."]])
    add_docx_table(one, [["Outcomes, Not Just Code"], ["We give you outcomes, not just code. Code is the mechanism. The product is operating lift: cleaner pipeline visibility, faster RFQ triage, fewer missed follow-ups, weekly executive decision rhythm, and reusable operating memory."]])
    add_docx_table(one, [["System Area", "Client-Facing Outcome"], ["Signal inbox", "Early project, permit, planning, ownership, franchise, broker, and RFQ signals captured in one place."], ["Qualification engine", "Pursue, watch, or pass recommendations based on Level-specific fit criteria."], ["Procore bridge", "Wrap around Procore first as the current source of truth, then reduce monolith dependence only where the owned layer proves stronger."], ["RFQ command center", "Scope summaries, deadlines, missing requirements, go/no-go memos, and bid tasks."], ["Executive brief", "Weekly view of what changed, what is stuck, and what leadership must decide."]])
    add_docx_table(one, [["Term Area", "Skeleton"], ["Engagement", "90-day pilot for Level Intelligence OS v1."], ["Deliverables", "Data model, Procore-first bridge, signal inbox, scoring framework, RFQ workflow, approval queue, weekly brief, success report, phase-two roadmap."], ["Ownership", "Level owns data, operating memory, domain model, workflow definitions, and exportable generated records."], ["Commercial", "Fixed-fee pilot, build plus retainer, strategy sprint, or long-term internal software partnership."]])
    one.save(OUT / "Level Intelligence OS - Official One Pager.docx")

    doc = Document()
    docx_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    add_title(doc, "Level Intelligence OS", "Official dossier")
    doc.add_paragraph("A private construction intelligence operating system for finding, qualifying, routing, and converting opportunities through Level's own operating judgment.")
    for heading, body in [
        ("Executive Summary", "Level Construction already has the hard part: trusted relationships, construction judgment, development instincts, hospitality experience, capital discipline, and a practical sense for which opportunities are worth time. The opportunity is to turn that judgment into a private operating system that Level owns."),
        ("The Business Problem", "Level's advantage is operational, but much of that advantage still lives in places that software cannot easily use: emails, RFQ packages, attachments, bid notes, permit comments, spreadsheets, project folders, meeting notes, calls, relationship memory, and executive judgment."),
        ("Outcomes, Not Just Code", "We give you outcomes, not just code. Code is the mechanism. The product is operating lift: cleaner pipeline visibility, faster triage, better RFQ discipline, fewer missed follow-ups, weekly executive rhythm, and reusable enterprise memory."),
    ]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    add_docx_table(doc, [["System Area", "Client-Facing Outcome"], ["Signal inbox", "Early project, permit, planning, ownership, franchise, broker, and RFQ signals captured in one place."], ["Opportunity graph", "Accounts, contacts, sites, jurisdictions, RFQs, bids, permits, and decisions connected."], ["Qualification engine", "Pursue, watch, or pass recommendations based on Level-specific fit criteria."], ["Relationship memory", "Searchable account/contact context, relationship owners, last touch, and next follow-up."], ["RFQ command center", "Scope summaries, deadlines, missing requirements, go/no-go memos, and bid tasks."], ["Procore bridge", "Procore remains the execution source of truth at first while Level Intelligence OS adds intelligence, pursuit, RFQ, approval, and executive workflow around it."], ["Approval queue", "Consequential actions routed to humans before external movement."], ["Executive brief", "Weekly view of what changed, what is stuck, what needs action, and what leadership must decide."]])
    doc.add_heading("Procore-First, Monolith-Later Strategy", level=1)
    doc.add_paragraph("The recommended first move is to wrap around Procore, not replace it. The long-term goal is to reduce reliance on monolithic software only after Level's owned operating layer proves which workflows can be handled with better fit, lower friction, and clear ROI.")
    add_docx_table(doc, [["Infrastructure Pattern", "Why It Matters"], ["Connected construction data model", "Accounts, contacts, sites, RFQs, bids, permits, approvals, decisions, and outcomes become one operating memory."], ["Human-in-the-loop controls", "Agents prepare, recommend, and route work while humans retain control over consequential actions."], ["Feedback loops", "Wins, losses, overrides, delays, and vendor outcomes improve future recommendations."]])
    doc.add_heading("90-Day Pilot", level=1)
    add_docx_table(doc, [["Period", "Focus", "Deliverable"], ["Weeks 1-2", "Foundation", "Data model, roles, active opportunity/account/contact views, executive dashboard skeleton."], ["Weeks 3-5", "Intelligence", "Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest."], ["Weeks 6-8", "Workflow", "RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking."], ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, outcome capture, adoption review, phase-two roadmap."]])
    doc.add_heading("Commercial And Ownership Principles", level=1)
    add_docx_table(doc, [["Area", "Recommended Position"], ["Engagement", "90-day pilot for a private, Level-specific construction intelligence operating system."], ["Ownership", "Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. Code ownership/licensing to be defined in final agreement."], ["Confidentiality", "Level data is confidential. No external reuse without written permission. Role-based access and audit logging required."], ["Commercial structure", "Fixed-fee pilot, build fee plus monthly operating retainer, strategy/design sprint followed by implementation, or long-term internal software partnership."]])
    doc.add_heading("Closing Position", level=1)
    doc.add_paragraph("The first version gives Level cleaner visibility and faster decisions. The strategic version gives Level a proprietary operating memory that improves how the company finds, qualifies, wins, and executes work.")
    doc.save(OUT / "Level Intelligence OS - Official Dossier.docx")

    packet = Document()
    docx_styles(packet)
    sec = packet.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    add_title(packet, "Level Intelligence OS", "Comprehensive Client Packet")
    packet.add_paragraph("A private construction intelligence operating system built from the ground up around Level's own relationships, standards, RFQs, approvals, markets, capital context, and operating judgment.")
    packet.add_paragraph("Commercial promise: We give you outcomes, not just code. Code is the mechanism. The product is operating lift.")

    packet.add_heading("1. Executive Thesis", level=1)
    packet.add_paragraph("Level Construction already has the hard part: trusted relationships, construction judgment, development instincts, hospitality experience, capital discipline, vendor memory, and practical pattern recognition about which opportunities are worth time. The opportunity is to convert that judgment into a private operating system that Level owns.")
    packet.add_paragraph("This is not a generic AI layer, chatbot, dashboard, CRM migration, or forced Procore replacement. It is an internal operating layer around business development, preconstruction, project pursuit, development, hospitality, capital workflows, and executive management.")

    packet.add_heading("2. The Pain Before The Platform", level=1)
    add_docx_table(packet, [
        ["Pain Point", "What It Looks Like", "Why It Matters"],
        ["Scattered signals", "Permits, planning activity, ownership changes, broker notes, franchise news, RFQs, and relationship tips arrive in different places.", "Promising opportunities can be seen late, duplicated, or missed."],
        ["Trapped judgment", "Fit, risk, timing, margin, and relationship context live in individual memory, inboxes, and informal notes.", "Level's advantage is real but hard to search, compare, hand off, or improve."],
        ["RFQ friction", "Deadlines, missing documents, addenda, site risk, scope ambiguity, and submission requirements are buried in language-heavy packages.", "Teams translate documents into tasks and may surface risk late."],
        ["Status burden", "Leadership needs to ask what changed, what is stuck, who owns it, and what needs approval.", "Executive time is spent assembling context instead of exercising judgment."],
    ])

    packet.add_heading("3. Built For Level, Not For The Category", level=1)
    packet.add_paragraph("Mercator and similar tools are useful category references because they show that earlier construction intelligence has value. But Level Intelligence OS is not a Mercator replacement pitch and not a vendor comparison exercise.")
    packet.add_paragraph("The strategic point is that Level should own software built from the ground up around Level's own relationships, standards, RFQ process, approval rules, development instincts, hospitality pipeline, capital context, and executive operating rhythm.")
    add_docx_table(packet, [
        ["Generic Vendor Logic", "Level Intelligence OS Logic"],
        ["Expose market or project signals.", "Expose signals and translate them into Level-specific opportunities, owners, scores, and next actions."],
        ["Use vendor-defined fields and workflows.", "Represent Level's own domain model: accounts, contacts, sites, RFQs, permits, bids, approvals, relationships, and decisions."],
        ["Optimize for broad category adoption.", "Optimize for Level's operating judgment, pursuit discipline, leadership rhythm, and long-term ownership."],
        ["Data remains mainly a tool output.", "Data becomes Level-owned operating memory that compounds over time."],
    ])

    packet.add_heading("4. Proven Infrastructure Patterns We Bring", level=1)
    packet.add_paragraph("The build can reuse internal construction-platform patterns already worked through in prior platform development: data models, workflow queues, approval gates, agent handoffs, document extraction, audit trails, dashboard surfaces, and feedback loops. The point is not to expose any prior build; the point is to lower risk and move faster.")
    add_docx_table(packet, [
        ["Pattern", "How It Applies To Level"],
        ["Construction domain model", "Accounts, contacts, sites, projects, permits, RFQs, bids, budget items, approvals, outcomes, and operating notes can be represented as connected records."],
        ["Workflow state machines", "Opportunities, RFQs, permits, approvals, and later payment or vendor workflows move through explicit statuses instead of scattered messages."],
        ["Specialist agent structure", "Orchestrated agents handle bounded jobs such as signal intake, qualification, RFQ extraction, permit monitoring, executive briefing, and finance context."],
        ["Human-in-the-loop controls", "Authority levels, approval queues, audit logs, escalation rules, and explicit human review keep consequential actions under Level control."],
        ["Document intelligence", "RFQs, permits, invoices, emails, addenda, and project narratives become deadlines, risks, tasks, missing information, and decision memos."],
        ["Feedback loops", "Wins, losses, overrides, delays, vendor outcomes, and cycle times become signals for better future recommendations."],
    ])

    packet.add_heading("5. Procore Bridge Strategy", level=1)
    packet.add_paragraph("Level Intelligence OS should wrap around Procore first, not replace it. Procore remains the execution source of truth while Level Intelligence OS becomes the intelligence, pursuit, RFQ, relationship, approval, and executive layer around it.")
    add_docx_table(packet, [
        ["Phase", "Position", "Outcome"],
        ["1. Read and reference", "Treat Procore as the current project source of truth. Pull or reference only the data needed for visibility and decisions.", "No disruption to field or project teams."],
        ["2. Workflow overlay", "Use Level Intelligence OS for signals, RFQs, go/no-go memos, relationship memory, approvals, and executive briefs.", "Level gains operating rhythm without forcing a platform migration."],
        ["3. Evidence-based rationalization", "Measure which workflows are better handled inside Level's own system and which should stay in Procore.", "Software decisions become based on observed value."],
        ["4. Selective replacement", "Move modules away from monolithic software only when the owned layer is clearly better, safer, and cheaper to operate.", "Dependency reduces over time without risky rip-and-replace."],
    ])

    packet.add_heading("6. Market Shift: Forward-Deployed Software", level=1)
    packet.add_paragraph("The software market is moving from fixed SaaS screens toward company-specific AI systems deployed inside real workflows. OpenAI's 2026 Deployment Company and forward-deployed engineering model are a clear market signal: enterprise value increasingly comes from fitting intelligence into real operating constraints, not simply licensing a generic tool.")
    add_docx_table(packet, [
        ["Shift", "Old Pattern", "New Pattern"],
        ["Static SaaS to internal systems", "The company adapts to vendor-defined screens and workflows.", "Software adapts to the company's language, judgment, data, approvals, and operating rhythm."],
        ["Generic automation to agentic work", "Rules handle narrow tasks but struggle with messy language-heavy work.", "Agents read, synthesize, route, draft, and learn from RFQs, emails, permits, notes, and decisions."],
        ["Vendor lock-in to owned memory", "Operating history lives inside external tools and exports.", "The company's decisions, outcomes, and workflows become a proprietary asset."],
        ["Implementation to forward deployment", "Software is installed and handed off.", "Experts map work, ship inside real constraints, tune with operators, and turn adoption into measurable outcomes."],
        ["Manual change requests to language-native change", "New workflows require slow product cycles or vendor roadmaps.", "Teams can describe additions, omissions, and changes in natural language, then convert them into owned software modules."],
    ])

    packet.add_heading("7. What The System Does", level=1)
    add_docx_table(packet, [
        ["Capability", "What It Does", "Client-Facing Outcome"],
        ["Find", "Detects public and private project signals before they become obvious.", "Level sees more opportunities earlier."],
        ["Qualify", "Scores opportunities using Level's criteria, history, relationships, and capacity.", "Leadership spends less time on poor-fit work."],
        ["Route", "Assigns owners, next actions, approvals, and follow-ups.", "Good opportunities do not stall in inboxes."],
        ["Convert", "Supports RFQs, go/no-go memos, pursuit briefs, and executive reporting.", "Level turns intelligence into action."],
        ["Learn", "Captures decisions, overrides, wins, losses, delays, and outcomes.", "Every cycle improves the next score, brief, and recommendation."],
    ])

    packet.add_heading("8. Product Modules", level=1)
    add_docx_table(packet, [
        ["Module", "What It Includes", "Why It Matters"],
        ["Signal inbox", "Permits, planning agendas, title/ownership changes, rezoning, franchise expansion, broker notes, inbound RFQs, CRM/email activity.", "Early signals become reviewable records instead of scattered noise."],
        ["Opportunity graph", "Accounts, contacts, sites, jurisdictions, RFQs, bids, permits, decisions, owners, next actions.", "Level sees relationships between people, projects, sites, and pursuits."],
        ["Qualification engine", "Pursue/watch/pass, confidence, reasoning, missing info, risks, next action, owner, comparable history.", "Level's own standards become repeatable without replacing human judgment."],
        ["Relationship memory", "Relationship owner, last touch, prior work, warm intros, next follow-up, leadership brief.", "Relationship context becomes searchable and actionable."],
        ["RFQ command center", "Scope summary, deadlines, addenda, required attachments, missing documents, risks, go/no-go memo, bid tasks.", "RFQs become decision-ready work."],
        ["Approval queue", "Recommendations, evidence, confidence, approver, decision, outcome.", "Agents prepare work while humans control consequential actions."],
        ["Executive brief", "Top signals, active pursuits, stalled items, RFQs due, approvals needed, relationship follow-ups, lessons learned.", "Leadership gets a weekly decision rhythm from live system data."],
    ])

    packet.add_heading("9. Agentic Operating Model", level=1)
    add_docx_table(packet, [
        ["Workflow Agent", "Role"],
        ["Signal Scout", "Finds, dedupes, summarizes, and routes project signals."],
        ["Qualification", "Scores opportunities against Level-specific criteria and explains the recommendation."],
        ["Relationship", "Builds account/contact briefs, tracks owners, suggests follow-ups, and supports warm intros."],
        ["RFQ/Bid", "Extracts requirements, deadlines, risks, missing documents, addenda, and go/no-go memos."],
        ["Permit/Entitlement", "Tracks jurisdiction status, review comments, days in review, and hidden timing risk."],
        ["Executive Briefing", "Turns system activity into a weekly decision brief."],
        ["Finance/Capital", "Connects opportunities to financing, underwriting, investor context, and capital constraints where relevant."],
    ])
    add_docx_table(packet, [
        ["Authority Level", "System Can Do", "Human Control"],
        ["Autonomous", "Classify signals, summarize RFQs, dedupe contacts, draft internal briefs.", "Logged for review."],
        ["Suggest", "Recommend owner, next action, pursue/watch/pass, follow-up draft.", "Human can accept, edit, or reject."],
        ["Approval Required", "External outreach, bid submission, vendor award, budget change, investor communication.", "Explicit approval required."],
        ["Manual", "Legal, contracts, financing, final strategic decisions.", "Human executes."],
    ])

    packet.add_heading("10. ROI And Value Model", level=1)
    add_docx_table(packet, [
        ["Value Lever", "How Value Is Created", "How To Measure"],
        ["Opportunity capture", "Signals become structured opportunities with owner, score, source, and next action.", "Signals captured, time to first action, conversion by source, wins by source."],
        ["Pursuit discipline", "Low-fit pursuits are filtered earlier while high-fit work receives faster attention.", "Go/no-go cycle time, win rate by score, avoided pursuit hours, margin by source."],
        ["Admin reduction", "Agents summarize, extract, draft, route, and remind across language-heavy workflows.", "Hours saved per RFQ, weekly report time saved, overdue task reduction."],
        ["Risk reduction", "Permits, missing documents, stalled decisions, and deadline risk surface sooner.", "Missed deadline rate, response time, unresolved blocker age."],
        ["Enterprise asset", "Relationships, decisions, outcomes, vendors, and market patterns become private memory.", "Record completeness, reuse of history, recommendation accuracy, override rate."],
    ])

    packet.add_heading("11. 90-Day Pilot", level=1)
    add_docx_table(packet, [
        ["Period", "Focus", "Client-Visible Deliverable"],
        ["Weeks 1-2", "Foundation", "Data model, roles, active opportunities, account/contact layer, executive dashboard skeleton."],
        ["Weeks 3-5", "Intelligence", "Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest."],
        ["Weeks 6-8", "Workflow", "RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking."],
        ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, outcome tracking, adoption review, phase-two roadmap."],
    ])

    packet.add_heading("12. Phase Two And System Expansion", level=1)
    add_docx_table(packet, [
        ["Expansion Area", "Potential Scope"],
        ["Permit and entitlement monitoring", "Jurisdiction status, comments, timing risk, follow-up drafts, days in review."],
        ["Email, calendar, CRM, storage integration", "Source access, relationship context, meeting history, follow-up capture."],
        ["Procore read-only bridge", "Treat Procore as the current project source of truth while Level Intelligence OS adds intelligence, pursuit, and executive workflow above it."],
        ["Vendor and subcontractor memory", "Performance history, issue patterns, pricing, reliability, project outcomes."],
        ["Finance and capital workflow", "Underwriting context, investor notes, financing constraints, capital approval support."],
        ["Software rationalization", "Usage audit to decide what should remain, integrate, downgrade, or eventually be replaced by Level-owned modules."],
    ])

    packet.add_heading("13. Commercial, Ownership, And Trust Principles", level=1)
    add_docx_table(packet, [
        ["Area", "Recommended Position"],
        ["Engagement", "90-day pilot to design, build, test, and operationalize Level Intelligence OS v1."],
        ["Pilot deliverables", "Workflow map, data model, opportunity/account/contact database, signal inbox, scoring framework, RFQ workflow, approval queue, weekly executive brief, success report, phase-two roadmap."],
        ["Ownership", "Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. Code ownership/licensing to be defined in the final agreement."],
        ["Model infrastructure", "AI model providers remain replaceable infrastructure, not the strategic asset."],
        ["Confidentiality", "Level data is confidential. No external reuse without written permission. Role-based access and audit logs are required."],
        ["Commercial options", "Fixed-fee pilot, build fee plus monthly operating retainer, strategy/design sprint followed by implementation, or long-term internal software partnership."],
    ])

    packet.add_heading("14. Discovery Agenda", level=1)
    add_docx_table(packet, [
        ["Discovery Topic", "Questions To Answer"],
        ["Opportunity sources", "Where do new opportunities originate today? Which sources matter most?"],
        ["Qualification criteria", "What makes a Level opportunity attractive or unattractive? Who decides?"],
        ["RFQ workflow", "Where do RFQs live? How are deadlines, addenda, tasks, and go/no-go decisions handled?"],
        ["Relationship memory", "Which accounts, contacts, owners, brands, developers, and vendors matter most?"],
        ["Executive reporting", "What does leadership need to see every week? What is currently hard to assemble?"],
        ["Systems and data", "Which tools are in use today: Procore, CRM, email, calendar, accounting, storage? What should be read, synced, imported, or left alone?"],
        ["Pilot success", "What would make the pilot obviously worth continuing?"],
    ])

    packet.add_heading("15. Closing Position", level=1)
    packet.add_paragraph("Level Intelligence OS should be presented as a serious operating asset, not an AI experiment. The first version gives Level one place to see opportunities, understand fit, track RFQs, preserve relationship memory, and route decisions.")
    packet.add_paragraph("The longer-term version becomes Level's private construction intelligence layer: a system that learns from Level's signals, relationships, decisions, wins, losses, and operating judgment.")
    packet.save(OUT / "Level Intelligence OS - Comprehensive Client Packet.docx")


SLIDES = [
    ("Level Intelligence OS", "A private construction intelligence operating system for finding, qualifying, routing, and converting opportunities through Level's own operating judgment.", ["Early signals", "Level-specific qualification", "Executive decision rhythm"]),
    ("Level's advantage is real. It is just scattered.", "The highest-value knowledge lives in RFQs, emails, calls, spreadsheets, permit threads, broker notes, project folders, and individual memory.", ["Active pursuits are hard to compare.", "RFQ risk is buried in documents.", "Relationship context depends on memory.", "Decisions are hard to learn from after the fact."]),
    ("The cost is missed operating leverage.", "Scattered judgment creates blind spots, duplicated work, slower decisions, weaker follow-up, and less learning from wins and losses.", ["Missed early opportunities", "Poor-fit pursuits consuming attention", "Leadership status chasing", "Past lessons not improving future scoring"]),
    ("Language-heavy work can finally become software.", "Agents can read, structure, route, and remember RFQs, emails, permit comments, meeting notes, relationship context, project narratives, and executive judgment.", ["RFQs become deadlines and go/no-go memos.", "Relationship notes become account memory.", "Permit comments become blockers and follow-ups.", "Decisions become reusable learning signals."]),
    ("This is not a vendor clone. It is built for Level.", "Tools like Mercator show that earlier construction intelligence has value. But the goal is not to compete with Mercator or any other vendor. The goal is to build software from the ground up around Level's operating model.", ["Generic tools expose signals.", "Level Intelligence OS encodes Level's relationships and standards.", "It supports Level's RFQ, approval, capital, and executive workflows.", "It becomes Level-owned operating memory."]),
    ("We bring proven platform primitives.", "Level gets the benefit of prior construction-platform learning without exposing any prior client work or making this feel generic.", ["Connected construction data model", "Workflow state machines", "Specialist agent handoffs", "Approval gates and audit logs", "Dashboards and executive briefs", "Feedback loops from outcomes"]),
    ("Procore is the bridge, not the battlefield.", "The first version should wrap around Procore as the current source of truth. Replacement only comes later, module by module, where the owned system proves stronger.", ["Read and reference Procore data", "Add intelligence above execution records", "Avoid disruption to active teams", "Measure friction and value", "Earn the right to replace"]),
    ("Software is becoming forward-deployed and company-specific.", "The market is shifting from static SaaS toward AI systems built inside real operating constraints by experts who map workflows, ship, tune, and measure outcomes.", ["OpenAI's Deployment Company is a market signal.", "Generic tools matter less than fit.", "Language-native systems can add and omit modules faster.", "Owned operating memory becomes a strategic asset."]),
    ("A private operating layer around Level's pursuit engine.", "One place to see the signal, understand the fit, route the work, and track the decision.", ["Find: detect project, permit, planning, ownership, franchise, broker, and RFQ signals.", "Qualify: score opportunities against Level's criteria.", "Route: assign owners, next actions, and approvals.", "Convert: support RFQs, go/no-go memos, pursuit briefs, and weekly executive review."]),
    ("We give you outcomes, not just code.", "Code is the mechanism. Operating outcomes are the product.", ["Pipeline visibility", "RFQ discipline", "Decision speed", "Executive rhythm", "Enterprise memory"]),
    ("Earlier signals, structured before they become noise.", "The signal inbox turns permits, planning agendas, title changes, franchise expansion, broker notes, inbound RFQs, CRM, and email activity into structured records.", ["Source", "Market", "Owner", "Confidence", "Score", "Next action"]),
    ("The score reflects Level's judgment.", "Each opportunity is scored against the criteria Level actually uses to decide whether work is worth pursuing.", ["Market fit", "Asset type", "Relationship strength", "Owner quality", "Entitlement complexity", "Margin potential", "Capacity", "Strategic value"]),
    ("Relationship context becomes operating memory.", "Accounts, contacts, internal relationship owners, prior projects, RFQs, touches, and next follow-ups become searchable and actionable.", ["Who do we know?", "Who owns the relationship?", "What have we built or bid?", "What should we say next?"]),
    ("Turn RFQs into decision-ready work.", "RFQs arrive as language-heavy packages. Level Intelligence OS extracts the operating facts and prepares the go/no-go decision.", ["Scope summary", "Deadline extraction", "Missing information", "Schedule and site risk", "Comparable past work", "Go/no-go memo"]),
    ("The system prepares. Level decides.", "Consequential actions remain under human control. The system drafts, recommends, routes, and logs.", ["Autonomous: summarize, classify, dedupe", "Suggest: owner, next action, pursue/watch/pass", "Approval required: outreach, bid submission, vendor award, budget change", "Manual: legal, financing, contracts, final strategic decisions"]),
    ("Leadership sees what changed, stuck, and needs judgment.", "The weekly executive brief replaces status gathering with decision support.", ["Top new signals", "Highest-scoring pursuits", "RFQs due soon", "Stalled items", "Pending approvals", "Relationship follow-ups", "Decisions needed"]),
    ("ROI is operational first, strategic second, compounding third.", "The bigger value is better pursuit discipline, faster decisions, fewer blind spots, and proprietary operating memory.", ["Operational ROI: less admin and faster RFQ review", "Strategic ROI: better selectivity and relationship leverage", "Enterprise ROI: Level-owned operating memory"]),
    ("A contained pilot with visible operating value.", "The 90-day pilot proves usefulness in workflows Level already runs.", ["Weeks 1-2: foundation", "Weeks 3-5: intelligence", "Weeks 6-8: workflow", "Weeks 9-12: operating rhythm"]),
    ("Build the system that makes Level sharper every week.", "The first version creates visibility and decision speed. The long-term version becomes Level's proprietary construction intelligence layer.", ["Level owns its data and generated operating memory.", "AI providers remain replaceable infrastructure.", "Export rights are preserved.", "Sensitive actions require approval.", "Approve a 90-day pilot."]),
]


def draw_slide(c, idx, title, subtitle, items):
    pw, ph = landscape(LETTER)
    c.setFillColor(PAPER)
    c.rect(0, 0, pw, ph, stroke=False, fill=True)
    c.setFillColor(DARK)
    c.rect(0, ph - 46, pw, 46, stroke=False, fill=True)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(0.42 * inch, ph - 28, "Level Intelligence OS")
    c.setFont("Helvetica", 8)
    c.drawRightString(pw - 0.42 * inch, ph - 28, f"{idx:02d} / {len(SLIDES):02d}")

    left = 0.55 * inch
    top = ph - 0.88 * inch
    c.setFillColor(ACCENT)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(left, top, "EXECUTIVE DECK")
    c.setFillColor(INK)
    c.setFont("Helvetica", 30 if len(title) < 48 else 25)
    y = top - 28
    for line in wrap_canvas(title, 29 if len(title) < 48 else 36):
        c.drawString(left, y, line)
        y -= 33 if len(title) < 48 else 28
    y -= 6
    c.setFillColor(MUTED)
    c.setFont("Helvetica", 11)
    for line in wrap_canvas(subtitle, 68):
        c.drawString(left, y, line)
        y -= 15

    panel_x = pw * 0.54
    panel_y = ph - 1.08 * inch
    panel_w = pw - panel_x - 0.55 * inch
    panel_h = ph - 1.58 * inch
    c.setFillColor(colors.white)
    c.setStrokeColor(LINE)
    c.rect(panel_x, panel_y - panel_h, panel_w, panel_h, stroke=True, fill=True)
    c.setStrokeColor(INK)
    c.line(panel_x + 16, panel_y - 34, panel_x + panel_w - 16, panel_y - 34)
    c.setFillColor(INK)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(panel_x + 16, panel_y - 22, "Operating view")

    row_y = panel_y - 58
    for i, item in enumerate(items[:8]):
        c.setFillColor(ACCENT if i % 3 == 0 else YELLOW if i % 3 == 1 else INK)
        c.rect(panel_x + 16, row_y - 5, 5, 5, stroke=False, fill=True)
        c.setFillColor(INK)
        c.setFont("Helvetica-Bold", 8.5)
        lines = wrap_canvas(item, 42)
        c.drawString(panel_x + 30, row_y, lines[0])
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 7.2)
        for more in lines[1:3]:
            row_y -= 10
            c.drawString(panel_x + 30, row_y, more)
        row_y -= 23

    c.setFillColor(DARK)
    c.rect(left, 0.55 * inch, pw * 0.42, 0.52 * inch, stroke=False, fill=True)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(left + 12, 0.84 * inch, "Find. Qualify. Route. Convert. Learn.")
    c.setFont("Helvetica", 7.5)
    c.setFillColor(colors.HexColor("#dce3e8"))
    c.drawString(left + 12, 0.66 * inch, "Outcomes, not just code.")


def wrap_canvas(text, chars):
    words = text.split()
    lines, line = [], ""
    for word in words:
        test = f"{line} {word}".strip()
        if len(test) <= chars:
            line = test
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def build_slides_pdf():
    c = canvas.Canvas(str(OUT / "Level Intelligence OS - Executive Slide Deck.pdf"), pagesize=landscape(LETTER))
    for i, (title, subtitle, items) in enumerate(SLIDES, start=1):
        draw_slide(c, i, title, subtitle, items)
        c.showPage()
    c.save()


def build_html_files():
    css = """
    :root{--ink:#121820;--muted:#5b6775;--line:#d8dde3;--paper:#f7f7f4;--surface:#fff;--accent:#0e806c;--dark:#17202b;--yellow:#c99b2e}
    *{box-sizing:border-box} body{margin:0;background:var(--paper);color:var(--ink);font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}
    a{color:inherit}.wrap{max-width:1180px;margin:auto;padding:34px}.top{display:flex;justify-content:space-between;gap:24px;border-bottom:2px solid var(--ink);padding-bottom:18px}.eyebrow{color:var(--accent);font-weight:850;font-size:12px;letter-spacing:.08em;text-transform:uppercase;margin:0 0 12px}h1{font-size:clamp(44px,7vw,88px);line-height:.9;margin:0;letter-spacing:0}h2{font-size:clamp(28px,4vw,52px);line-height:1;margin:0 0 18px}p{color:var(--muted);line-height:1.55;font-size:17px}.grid{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}.section{padding:54px 0;border-bottom:1px solid var(--line)}.panel{border-top:2px solid var(--ink);padding-top:13px}.panel b{display:block;margin-bottom:8px}.panel span{display:block;color:var(--muted);line-height:1.4}.dark{background:var(--dark);color:white;margin:0 calc(50% - 50vw);padding:64px calc(50vw - 50%)}.dark p,.dark span{color:rgba(255,255,255,.75)}table{width:100%;border-collapse:collapse;background:white}th{background:var(--dark);color:white;text-align:left}th,td{border:1px solid var(--line);padding:12px;vertical-align:top}.cta{display:flex;gap:12px;flex-wrap:wrap}.btn{background:var(--dark);color:white;text-decoration:none;padding:12px 16px;font-weight:800}.btn.alt{background:white;color:var(--ink);border:1px solid var(--line)}@media(max-width:820px){.top,.grid{grid-template-columns:1fr;display:grid}.wrap{padding:22px}}
    """

    one_body = """
    <main class="wrap">
      <header class="top"><div><p class="eyebrow">Official one-page dossier</p><h1>Level Intelligence OS</h1></div><p>A private construction intelligence operating system for finding, qualifying, routing, and converting opportunities through Level's own operating judgment.</p></header>
      <section class="section"><h2>Level's advantage is real. It is just scattered.</h2><div class="grid"><div class="panel"><b>Scattered signals</b><span>Permits, planning agendas, broker notes, franchise expansion, RFQs, and relationship updates arrive in different places.</span></div><div class="panel"><b>Trapped judgment</b><span>Fit, risk, timing, margin, and relationship context live in memory, inboxes, and informal notes.</span></div><div class="panel"><b>Status burden</b><span>Leadership needs what changed, what is stuck, what needs approval, and who owns the next move.</span></div></div></section>
      <section class="section dark"><h2>We give you outcomes, not just code.</h2><p>Code is the mechanism. The product is operating lift: cleaner pipeline visibility, faster RFQ triage, fewer missed follow-ups, weekly executive decision rhythm, and reusable operating memory.</p></section>
      <section class="section"><h2>What Level gets</h2><table><tr><th>System Area</th><th>Client-Facing Outcome</th></tr><tr><td>Signal inbox</td><td>Early project, permit, planning, ownership, franchise, broker, and RFQ signals captured in one place.</td></tr><tr><td>Qualification engine</td><td>Pursue, watch, or pass recommendations based on Level-specific fit criteria.</td></tr><tr><td>Procore bridge</td><td>Wrap around Procore first as the current source of truth, then reduce monolith dependence only where the owned layer proves stronger.</td></tr><tr><td>RFQ command center</td><td>Scope summaries, deadlines, missing requirements, go/no-go memos, and bid tasks.</td></tr><tr><td>Executive brief</td><td>Weekly view of what changed, what is stuck, and what leadership must decide.</td></tr></table></section>
      <section class="section"><h2>90-day pilot</h2><table><tr><th>Period</th><th>Focus</th><th>Deliverable</th></tr><tr><td>Weeks 1-2</td><td>Foundation</td><td>Data model, roles, active pursuits, contacts, executive view.</td></tr><tr><td>Weeks 3-5</td><td>Intelligence</td><td>Signal inbox, summaries, scoring, relationship briefs.</td></tr><tr><td>Weeks 6-8</td><td>Workflow</td><td>RFQ command center, go/no-go memo, approval queue.</td></tr><tr><td>Weeks 9-12</td><td>Rhythm</td><td>Weekly brief, outcome capture, adoption review, roadmap.</td></tr></table></section>
    </main>
    """
    write_html("index.html", "Level Intelligence OS - One Page", css, one_body)

    dossier_body = """
    <main class="wrap">
      <header class="top"><div><p class="eyebrow">Official dossier</p><h1>Level Intelligence OS</h1></div><p>A private operating layer around Level's business development, preconstruction, project pursuit, development, hospitality, and capital workflows.</p></header>
      <section class="section"><h2>Executive summary</h2><p>Level already has the hard part: trusted relationships, commercial construction judgment, development instincts, hospitality experience, capital discipline, and a practical sense for which opportunities are worth time. The opportunity is to turn that judgment into a private operating system that Level owns.</p><p><b>This is not built to compete with Mercator or clone any vendor. It is built from the ground up around Level's own relationships, standards, RFQs, approvals, markets, capital context, and operating judgment.</b></p></section>
      <section class="section"><h2>The business problem</h2><p>Level's advantage is operational, but much of that advantage lives in emails, RFQ packages, attachments, bid notes, permit comments, spreadsheets, project folders, meeting notes, calls, relationship memory, and executive judgment.</p></section>
      <section class="section dark"><h2>Outcomes, not just code</h2><p>The deliverable is a working operating rhythm: visible pipeline, better RFQ discipline, fewer missed follow-ups, approval rails, weekly executive brief, and compounding operating memory.</p></section>
      <section class="section"><h2>Procore-first, monolith-later</h2><p>The first move is to wrap around Procore, not replace it. Procore can remain the execution source of truth while Level Intelligence OS adds intelligence, pursuit, RFQ, approval, relationship, and executive workflow around it. Over time, Level can move modules away from monolithic software only where the owned layer proves better fit, lower friction, and clear ROI.</p></section>
      <section class="section"><h2>Operating system areas</h2><table><tr><th>Area</th><th>Outcome</th></tr><tr><td>Signal inbox</td><td>Early signals captured and routed.</td></tr><tr><td>Opportunity graph</td><td>Accounts, contacts, sites, RFQs, bids, permits, and decisions connected.</td></tr><tr><td>Qualification engine</td><td>Level-specific pursue, watch, or pass recommendations.</td></tr><tr><td>Relationship memory</td><td>Account and contact context becomes searchable.</td></tr><tr><td>RFQ command center</td><td>Decision-ready RFQ workflow.</td></tr><tr><td>Executive brief</td><td>Weekly view of changes, stuck items, and decisions needed.</td></tr></table></section>
      <section class="section"><h2>Pilot and ownership</h2><p>The 90-day pilot proves value in existing workflows. Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. AI model providers remain replaceable infrastructure.</p></section>
    </main>
    """
    write_html("dossier.html", "Level Intelligence OS - Dossier", css, dossier_body)

    packet_body = """
    <main class="wrap">
      <header class="top"><div><p class="eyebrow">Comprehensive client packet</p><h1>Level Intelligence OS</h1></div><p>A private construction intelligence operating system built from the ground up around Level's own relationships, standards, RFQs, approvals, markets, capital context, and operating judgment.</p></header>
      <section class="section dark"><h2>We give you outcomes, not just code.</h2><p>Code is the mechanism. The product is operating lift: cleaner opportunity visibility, faster RFQ triage, stronger go/no-go discipline, fewer missed follow-ups, a weekly executive decision rhythm, and Level-owned operating memory that compounds over time.</p></section>
      <section class="section"><h2>Executive thesis</h2><p>Level already has trusted relationships, construction judgment, development instincts, hospitality experience, capital discipline, vendor memory, and practical pattern recognition about which opportunities are worth time. The opportunity is to convert that judgment into a private operating system that Level owns.</p><p>The first version creates visibility and decision speed. The strategic version becomes Level's proprietary construction intelligence layer.</p></section>
      <section class="section"><h2>The pain before the platform</h2><div class="grid"><div class="panel"><b>Scattered signals</b><span>Permits, planning activity, ownership changes, broker notes, franchise news, RFQs, and relationship tips arrive in different places.</span></div><div class="panel"><b>Trapped judgment</b><span>Fit, risk, timing, margin, and relationship context live in individual memory, inboxes, and informal notes.</span></div><div class="panel"><b>RFQ friction</b><span>Deadlines, missing documents, addenda, site risk, scope ambiguity, and submission requirements are buried in language-heavy packages.</span></div></div></section>
      <section class="section"><h2>Built for Level, not for the category</h2><p>Mercator and similar tools show that earlier construction intelligence has value. But this is not a vendor clone, replacement pitch, or comparison exercise. Level Intelligence OS is built from the ground up around Level's relationships, standards, RFQ process, approval rules, development instincts, hospitality pipeline, capital context, and executive operating rhythm.</p><table><tr><th>Generic vendor logic</th><th>Level Intelligence OS logic</th></tr><tr><td>Expose market or project signals.</td><td>Translate signals into Level-specific opportunities, owners, scores, and next actions.</td></tr><tr><td>Use vendor-defined fields and workflows.</td><td>Represent Level's own domain model: accounts, contacts, sites, RFQs, permits, bids, approvals, relationships, and decisions.</td></tr><tr><td>Optimize for broad category adoption.</td><td>Optimize for Level's operating judgment, pursuit discipline, leadership rhythm, and long-term ownership.</td></tr></table></section>
      <section class="section"><h2>Proven infrastructure patterns</h2><p>The build can reuse internal construction-platform patterns already worked through in prior platform development: data models, workflow queues, approval gates, agent handoffs, document extraction, audit trails, dashboard surfaces, and feedback loops. The point is not to expose any prior build. The point is to lower risk and move faster.</p><table><tr><th>Pattern</th><th>How it applies to Level</th></tr><tr><td>Construction domain model</td><td>Accounts, contacts, sites, projects, permits, RFQs, bids, budget items, approvals, outcomes, and operating notes become connected records.</td></tr><tr><td>Workflow state machines</td><td>Opportunities, RFQs, permits, approvals, and later vendor workflows move through explicit statuses instead of scattered messages.</td></tr><tr><td>Specialist agent structure</td><td>Bounded agents handle signal intake, qualification, RFQ extraction, permit monitoring, executive briefing, and finance context.</td></tr><tr><td>Human controls</td><td>Authority levels, approval queues, audit logs, escalation rules, and human review keep consequential actions under Level control.</td></tr><tr><td>Feedback loops</td><td>Wins, losses, overrides, delays, vendor outcomes, and cycle times improve future recommendations.</td></tr></table></section>
      <section class="section"><h2>Procore bridge strategy</h2><p>Level Intelligence OS should wrap around Procore first, not replace it. Procore remains the execution source of truth while Level Intelligence OS becomes the intelligence, pursuit, RFQ, relationship, approval, and executive layer around it.</p><table><tr><th>Phase</th><th>Position</th><th>Outcome</th></tr><tr><td>Read and reference</td><td>Pull or reference only the Procore data needed for visibility and decisions.</td><td>No disruption to field or project teams.</td></tr><tr><td>Workflow overlay</td><td>Use Level Intelligence OS for signals, RFQs, go/no-go memos, relationship memory, approvals, and executive briefs.</td><td>Operating rhythm improves without forcing migration.</td></tr><tr><td>Evidence-based rationalization</td><td>Measure which workflows are better handled inside Level's own system and which should stay in Procore.</td><td>Software decisions become based on observed value.</td></tr><tr><td>Selective replacement</td><td>Move modules away from monolithic software only when the owned layer is clearly better, safer, and cheaper to operate.</td><td>Dependency reduces over time without risky rip-and-replace.</td></tr></table></section>
      <section class="section"><h2>Market shift</h2><p>The software market is moving from fixed SaaS screens toward company-specific AI systems deployed inside real workflows. OpenAI's 2026 Deployment Company and forward-deployed engineering model are a clear market signal: enterprise value increasingly comes from fitting intelligence into real operating constraints, not simply licensing a generic tool.</p><div class="grid"><div class="panel"><b>Static SaaS to internal systems</b><span>Software adapts to the company's language, judgment, data, approvals, and operating rhythm.</span></div><div class="panel"><b>Generic automation to agentic work</b><span>Agents can read, synthesize, route, draft, and learn from RFQs, emails, permits, notes, and decisions.</span></div><div class="panel"><b>Vendor lock-in to owned memory</b><span>Decisions, outcomes, and workflows become a proprietary company asset.</span></div></div></section>
      <section class="section"><h2>Operating loop</h2><table><tr><th>Capability</th><th>What it does</th><th>Outcome</th></tr><tr><td>Find</td><td>Detects public and private project signals.</td><td>Level sees more opportunities earlier.</td></tr><tr><td>Qualify</td><td>Scores opportunities using Level's criteria, history, relationships, and capacity.</td><td>Leadership spends less time on poor-fit work.</td></tr><tr><td>Route</td><td>Assigns owners, next actions, approvals, and follow-ups.</td><td>Good opportunities do not stall in inboxes.</td></tr><tr><td>Convert</td><td>Supports RFQs, go/no-go memos, pursuit briefs, and executive reporting.</td><td>Level turns intelligence into action.</td></tr><tr><td>Learn</td><td>Captures decisions, overrides, wins, losses, delays, and outcomes.</td><td>Every cycle improves the next recommendation.</td></tr></table></section>
      <section class="section"><h2>Product modules</h2><table><tr><th>Module</th><th>What it includes</th><th>Why it matters</th></tr><tr><td>Signal inbox</td><td>Permits, planning agendas, ownership changes, rezoning, franchise expansion, broker notes, inbound RFQs, CRM/email activity.</td><td>Early signals become reviewable records.</td></tr><tr><td>Qualification engine</td><td>Pursue/watch/pass, confidence, reasoning, missing info, risks, next action, owner, comparable history.</td><td>Level's standards become repeatable.</td></tr><tr><td>RFQ command center</td><td>Scope summary, deadlines, addenda, required attachments, missing documents, risks, go/no-go memo, bid tasks.</td><td>RFQs become decision-ready work.</td></tr><tr><td>Executive brief</td><td>Top signals, active pursuits, stalled items, RFQs due, approvals needed, relationship follow-ups, lessons learned.</td><td>Leadership gets a weekly decision rhythm.</td></tr></table></section>
      <section class="section"><h2>Agentic operating model</h2><table><tr><th>Workflow agent</th><th>Role</th></tr><tr><td>Signal Scout</td><td>Finds, dedupes, summarizes, and routes project signals.</td></tr><tr><td>Qualification</td><td>Scores opportunities against Level-specific criteria and explains the recommendation.</td></tr><tr><td>Relationship</td><td>Builds account/contact briefs, tracks owners, suggests follow-ups, and supports warm intros.</td></tr><tr><td>RFQ/Bid</td><td>Extracts requirements, deadlines, risks, missing documents, addenda, and go/no-go memos.</td></tr><tr><td>Executive Briefing</td><td>Turns system activity into a weekly decision brief.</td></tr></table></section>
      <section class="section"><h2>ROI and value model</h2><table><tr><th>Value lever</th><th>How value is created</th><th>How to measure</th></tr><tr><td>Opportunity capture</td><td>Signals become structured opportunities with owner, score, source, and next action.</td><td>Signals captured, time to first action, conversion by source.</td></tr><tr><td>Pursuit discipline</td><td>Low-fit pursuits are filtered earlier while high-fit work gets faster attention.</td><td>Go/no-go cycle time, win rate by score, avoided pursuit hours.</td></tr><tr><td>Risk reduction</td><td>Permits, missing documents, stalled decisions, and deadline risk surface sooner.</td><td>Missed deadline rate, response time, unresolved blocker age.</td></tr><tr><td>Enterprise asset</td><td>Relationships, decisions, outcomes, vendors, and market patterns become private memory.</td><td>Record completeness, reuse of history, recommendation accuracy.</td></tr></table></section>
      <section class="section"><h2>90-day pilot</h2><table><tr><th>Period</th><th>Focus</th><th>Client-visible deliverable</th></tr><tr><td>Weeks 1-2</td><td>Foundation</td><td>Data model, roles, active opportunities, account/contact layer, executive dashboard skeleton.</td></tr><tr><td>Weeks 3-5</td><td>Intelligence</td><td>Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest.</td></tr><tr><td>Weeks 6-8</td><td>Workflow</td><td>RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking.</td></tr><tr><td>Weeks 9-12</td><td>Operating rhythm</td><td>Weekly executive brief, outcome tracking, adoption review, phase-two roadmap.</td></tr></table></section>
      <section class="section"><h2>Commercial, ownership, and trust</h2><table><tr><th>Area</th><th>Recommended position</th></tr><tr><td>Engagement</td><td>90-day pilot to design, build, test, and operationalize Level Intelligence OS v1.</td></tr><tr><td>Ownership</td><td>Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records.</td></tr><tr><td>Model infrastructure</td><td>AI model providers remain replaceable infrastructure, not the strategic asset.</td></tr><tr><td>Confidentiality</td><td>Level data is confidential. No external reuse without written permission. Role-based access and audit logs are required.</td></tr></table></section>
      <section class="section"><h2>Discovery agenda</h2><p>Recommended next step: run a 60-minute workflow discovery session, collect sample RFQs and current active pursuits, then finalize the pilot scope and commercial structure.</p><div class="grid"><div class="panel"><b>Inputs</b><span>Active pursuits, sample RFQs, account/contact list, current tools, reporting format.</span></div><div class="panel"><b>Decisions</b><span>Qualification criteria, approval rules, pilot data sources, week-one build scope.</span></div><div class="panel"><b>Output</b><span>Final pilot scope, success metrics, timeline, and commercial structure.</span></div></div></section>
    </main>
    """
    write_html("comprehensive-packet.html", "Level Intelligence OS - Comprehensive Client Packet", css, packet_body)

    slide_css = css + """
    body{background:#111820}.deck{height:100svh;overflow-y:scroll;scroll-snap-type:y mandatory}.slide{height:100svh;scroll-snap-align:start;background:var(--paper);display:grid;grid-template-columns:1.05fr .95fr;gap:48px;padding:64px}.slide:nth-child(10),.slide:last-child{background:var(--dark);color:white}.slide:nth-child(10) p,.slide:nth-child(10) span,.slide:last-child p,.slide:last-child span{color:rgba(255,255,255,.78)}.num{color:var(--accent);font-weight:900;letter-spacing:.08em}.slide h2{font-size:clamp(42px,6vw,80px)}.mock{background:white;border:1px solid var(--line);padding:24px;align-self:center}.slide:nth-child(10) .mock,.slide:last-child .mock{background:rgba(255,255,255,.08);border-color:rgba(255,255,255,.2)}.item{border-top:2px solid var(--ink);padding:12px 0;color:var(--muted)}.slide:nth-child(10) .item,.slide:last-child .item{border-color:white;color:rgba(255,255,255,.82)}@media(max-width:900px){.slide{height:auto;min-height:100svh;grid-template-columns:1fr;padding:34px}.slide h2{font-size:44px}}
    """
    slides = ["<main class='deck'>"]
    for i, (title, subtitle, items) in enumerate(SLIDES, start=1):
        slides.append(f"<section class='slide'><div><p class='num'>{i:02d} / {len(SLIDES):02d}</p><h2>{escape(title)}</h2><p>{escape(subtitle)}</p></div><div class='mock'>")
        for item in items:
            slides.append(f"<div class='item'><b>{escape(item)}</b></div>")
        slides.append("</div></section>")
    slides.append("</main>")
    write_html("slides.html", "Level Intelligence OS - Executive Slide Deck", slide_css, "\n".join(slides))

    index = """
    <main class="wrap">
      <header class="top"><div><p class="eyebrow">Official package</p><h1>Level Intelligence OS</h1></div><p>Client-ready documents, HTML pages, and slide deck exports.</p></header>
      <section class="section"><h2>Files</h2><div class="grid"><a class="btn" href="index.html">One-page HTML</a><a class="btn" href="dossier.html">Dossier HTML</a><a class="btn" href="comprehensive-packet.html">Comprehensive HTML</a><a class="btn" href="slides.html">Slide deck HTML</a><a class="btn alt" href="Level%20Intelligence%20OS%20-%20Official%20One%20Pager.pdf">One-page PDF</a><a class="btn alt" href="Level%20Intelligence%20OS%20-%20Official%20Dossier.pdf">Dossier PDF</a><a class="btn alt" href="Level%20Intelligence%20OS%20-%20Comprehensive%20Client%20Packet.pdf">Comprehensive PDF</a><a class="btn alt" href="Level%20Intelligence%20OS%20-%20Executive%20Slide%20Deck.pdf">Slide deck PDF</a></div></section>
    </main>
    """
    write_html("package.html", "Level Intelligence OS - Official Package", css, index)


def write_html(filename, title, css, body):
    html = f"<!doctype html><html lang='en'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>{escape(title)}</title><style>{css}</style></head><body>{body}</body></html>"
    (OUT / filename).write_text(html, encoding="utf-8")


def main():
    ensure_dirs()
    build_one_page_pdf()
    build_dossier_pdf()
    build_comprehensive_packet_pdf()
    build_docx_files()
    build_slides_pdf()
    build_html_files()
    print(f"Official package generated in {OUT}")


if __name__ == "__main__":
    main()
