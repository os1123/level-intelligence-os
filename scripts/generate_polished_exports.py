from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Flowable,
    Frame,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/omar/Grounded/Projects/Level Intelligence OS")
OUT = ROOT / "exports"
OUT.mkdir(exist_ok=True)

INK = colors.HexColor("#11161d")
MUTED = colors.HexColor("#5f6b76")
LINE = colors.HexColor("#d9dee4")
PAPER = colors.HexColor("#f6f7f4")
SURFACE = colors.white
ACCENT = colors.HexColor("#0d806c")
DARK = colors.HexColor("#17202b")
GOLD = colors.HexColor("#bd8427")
BLUE = colors.HexColor("#4269d6")
PURPLE = colors.HexColor("#8b5ac2")
RISK = colors.HexColor("#b84a4a")


def pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "eyebrow": ParagraphStyle(
            "eyebrow",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10,
            textColor=ACCENT,
            spaceAfter=8,
            uppercase=True,
        ),
        "title": ParagraphStyle(
            "title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=44,
            leading=44,
            textColor=INK,
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=27,
            textColor=INK,
            spaceBefore=4,
            spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=INK,
            spaceBefore=10,
            spaceAfter=7,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=MUTED,
            spaceAfter=6,
        ),
        "body_dark": ParagraphStyle(
            "body_dark",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#dce3e8"),
            spaceAfter=6,
        ),
        "lead": ParagraphStyle(
            "lead",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=12,
            leading=17,
            textColor=MUTED,
            spaceAfter=12,
        ),
        "small": ParagraphStyle(
            "small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=MUTED,
        ),
        "cell_head": ParagraphStyle(
            "cell_head",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=10,
            textColor=INK,
        ),
        "cell": ParagraphStyle(
            "cell",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=MUTED,
        ),
        "white_head": ParagraphStyle(
            "white_head",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=11,
            textColor=colors.white,
        ),
        "white_cell": ParagraphStyle(
            "white_cell",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#dce3e8"),
        ),
    }
    return styles


S = pdf_styles()


def p(text, style="body"):
    return Paragraph(text, S[style])


def bullet_items(items, style="body"):
    return [p(f"- {item}", style) for item in items]


class Rule(Flowable):
    def __init__(self, color=LINE, width=1):
        super().__init__()
        self.color = color
        self.width = width
        self.height = 0.08 * inch

    def wrap(self, availWidth, availHeight):
        self.availWidth = availWidth
        return availWidth, self.height

    def draw(self):
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.width)
        self.canv.line(0, self.height / 2, self.availWidth, self.height / 2)


def make_table(rows, col_widths, header=True, dark_header=True):
    data = []
    for r_i, row in enumerate(rows):
        row_data = []
        for cell in row:
            if r_i == 0 and header:
                row_data.append(p(str(cell), "white_head" if dark_header else "cell_head"))
            else:
                row_data.append(p(str(cell), "cell"))
        data.append(row_data)
    tbl = Table(data, colWidths=col_widths, hAlign="LEFT")
    style = [
        ("BOX", (0, 0), (-1, -1), 0.6, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
    ]
    if header:
        style += [
            ("BACKGROUND", (0, 0), (-1, 0), DARK if dark_header else colors.HexColor("#f0f3f1")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white if dark_header else INK),
        ]
    tbl.setStyle(TableStyle(style))
    return tbl


def on_page(canv, doc):
    canv.saveState()
    w, h = doc.pagesize
    canv.setStrokeColor(LINE)
    canv.line(doc.leftMargin, h - 0.45 * inch, w - doc.rightMargin, h - 0.45 * inch)
    canv.setFont("Helvetica-Bold", 8)
    canv.setFillColor(INK)
    canv.drawString(doc.leftMargin, h - 0.32 * inch, "Level Intelligence OS")
    canv.setFont("Helvetica", 8)
    canv.setFillColor(MUTED)
    canv.drawRightString(w - doc.rightMargin, h - 0.32 * inch, f"Page {doc.page}")
    canv.restoreState()


def draw_paragraph(c, text, x, y, w, h, style):
    frame = Frame(x, y, w, h, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, showBoundary=0)
    frame.addFromList([Paragraph(text, style)], c)


def draw_wrapped(c, text, x, y_top, w, style, max_h=1.0 * inch):
    para = Paragraph(text, style)
    _, h = para.wrap(w, max_h)
    para.drawOn(c, x, y_top - h)
    return y_top - h


def one_page_pdf():
    path = OUT / "Level Intelligence OS - One Page Sheet.pdf"
    c = canvas.Canvas(str(path), pagesize=landscape(LETTER))
    pw, ph = landscape(LETTER)
    m = 28
    cw = pw - 2 * m
    y = ph - m

    def text_lines(text, font, size, max_w):
        words = text.split()
        lines, cur = [], ""
        for word in words:
            trial = (cur + " " + word).strip()
            if c.stringWidth(trial, font, size) <= max_w:
                cur = trial
            else:
                if cur:
                    lines.append(cur)
                cur = word
        if cur:
            lines.append(cur)
        return lines

    def draw_wrapped_text(text, x, top, max_w, font="Helvetica", size=7.2, leading=9, color=MUTED, max_lines=None):
        c.setFont(font, size)
        c.setFillColor(color)
        lines = text_lines(text, font, size, max_w)
        if max_lines:
            lines = lines[:max_lines]
        yy = top
        for line in lines:
            c.drawString(x, yy, line)
            yy -= leading
        return yy

    def box(x, top, w, h, fill=colors.white, stroke=LINE):
        c.setFillColor(fill)
        c.setStrokeColor(stroke)
        c.rect(x, top - h, w, h, fill=True, stroke=True)

    def heading(text, x, top, size=12, color=INK):
        c.setFillColor(color)
        c.setFont("Helvetica-Bold", size)
        c.drawString(x, top, text)

    def small_head(text, x, top):
        c.setFillColor(INK)
        c.setFont("Helvetica-Bold", 7.2)
        c.drawString(x, top, text)

    c.setFillColor(PAPER)
    c.rect(0, 0, pw, ph, fill=True, stroke=False)

    # Header
    c.setFillColor(ACCENT)
    c.setFont("Helvetica-Bold", 7.8)
    c.drawString(m, y, "CLIENT ONE-PAGE DOSSIER")
    y -= 31
    c.setFillColor(INK)
    c.setFont("Helvetica-Bold", 34)
    c.drawString(m, y, "Level Intelligence OS")
    c.setStrokeColor(INK)
    c.setLineWidth(1.2)
    c.line(m, y - 14, pw - m, y - 14)
    c.rect(pw - m - 42, y - 4, 34, 34, stroke=True, fill=False)
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(pw - m - 25, y + 7, "LX")
    y -= 32

    # Pain points
    pain_h = 58
    box(m, y, cw, pain_h)
    heading("The pain points", m + 10, y - 16, 11)
    col_gap = 12
    col_w = (cw - 20 - 2 * col_gap) / 3
    pain = [
        ("Signals are scattered", "Opportunities start in permits, planning, RFQs, brokers, brands, and relationships."),
        ("Judgment is trapped", "Context lives in people, inboxes, spreadsheets, and project memory."),
        ("Leadership chases status", "Executives need exceptions, decisions, risks, and next actions in one place."),
    ]
    for i, (h, b) in enumerate(pain):
        x = m + 10 + i * (col_w + col_gap)
        c.setStrokeColor(INK)
        c.line(x, y - 27, x + col_w, y - 27)
        small_head(h, x, y - 38)
        draw_wrapped_text(b, x, y - 48, col_w, size=6.6, leading=7.7, max_lines=2)
    y -= pain_h + 8

    # Outcome promise and feasibility side-by-side
    row_h = 82
    left_w = cw * 0.48
    right_w = cw - left_w - 10
    box(m, y, left_w, row_h, fill=DARK, stroke=DARK)
    heading("Outcomes, not just code", m + 10, y - 17, 11, colors.white)
    draw_wrapped_text(
        "Code is the mechanism. The product is a working operating rhythm: visible pipeline, cleaner RFQ decisions, approval rails, weekly executive briefs, and compounding operating memory.",
        m + 10,
        y - 34,
        left_w - 20,
        size=6.7,
        leading=8,
        color=colors.HexColor("#dce3e8"),
        max_lines=5,
    )
    x2 = m + left_w + 10
    box(x2, y, right_w, row_h)
    heading("Feasibility", x2 + 10, y - 17, 11)
    feas = [
        ("High feasibility", "Build above current tools first."),
        ("Main risk", "Data quality and adoption."),
        ("Best wedge", "Pipeline, RFQs, briefs."),
    ]
    f_w = (right_w - 20 - 2 * col_gap) / 3
    for i, (h, b) in enumerate(feas):
        xx = x2 + 10 + i * (f_w + col_gap)
        c.setStrokeColor(INK)
        c.line(xx, y - 29, xx + f_w, y - 29)
        small_head(h, xx, y - 41)
        draw_wrapped_text(b, xx, y - 52, f_w, size=6.6, leading=7.7, max_lines=2)
    y -= row_h + 8

    # What Level gets / why
    mid_h = 128
    mid_w = (cw - 12) / 2
    boxes = [
        ("What Level gets", [
            "Earlier signal detection across permits, planning, title changes, development news, broker notes, franchise expansion, and RFQs.",
            "Level-specific qualification scoring across geography, relationship strength, margin potential, entitlement risk, financing fit, and strategic value.",
            "Agent workflows for scouting, relationship briefs, RFQ summaries, go/no-go memos, permit follow-ups, executive briefs, and handoffs.",
            "Human approval rails for outreach, bid submission, vendor award, budget change, investor communication, and legal decisions.",
        ]),
        ("Why this matters now", [
            "Software is becoming internal and company-specific again.",
            "Language is now a software interface, so messy emails, RFQs, permit comments, and notes become structured work.",
            "Level can add, omit, and reshape workflows faster when it owns the operating layer.",
            "Every decision and outcome becomes training material for better future recommendations.",
        ]),
    ]
    for i, (h, items) in enumerate(boxes):
        x = m + i * (mid_w + 12)
        box(x, y, mid_w, mid_h)
        heading(h, x + 10, y - 17, 11)
        yy = y - 34
        for item in items:
            c.setFillColor(ACCENT)
            c.rect(x + 10, yy - 3, 3, 3, fill=True, stroke=False)
            yy = draw_wrapped_text(item, x + 18, yy, mid_w - 30, size=6.8, leading=7.7, max_lines=2) - 2
    y -= mid_h + 8

    # Value props and loop
    value_h = 74
    box(m, y, cw, value_h)
    heading("Value props and compounding loop", m + 10, y - 17, 11)
    vals = [
        ("Faster pursuit", "Signals become opportunities with owners, scores, and next actions."),
        ("Better margin", "Go/no-go discipline reduces wasted pursuits and highlights risk earlier."),
        ("Lower admin load", "Agents summarize, extract, draft, route, and remind across language-heavy work."),
        ("Outcome delivery", "Success is measured by operating lift, not by handing over a codebase."),
        ("Improve", "Every outcome sharpens the next score, brief, alert, and recommendation."),
    ]
    v_w = (cw - 20 - 4 * 8) / 5
    for i, (h, b) in enumerate(vals):
        x = m + 10 + i * (v_w + 8)
        c.setStrokeColor(INK)
        c.line(x, y - 29, x + v_w, y - 29)
        small_head(h, x, y - 40)
        draw_wrapped_text(b, x, y - 51, v_w, size=6.4, leading=7.3, max_lines=3)
    y -= value_h + 8

    # Terms
    term_h = 98
    box(m, y, cw, term_h, fill=DARK, stroke=DARK)
    heading("Term sheet skeleton", m + 10, y - 17, 11, colors.white)
    terms = [
        ("Engagement", "90-day pilot to design and build Level Intelligence OS v1."),
        ("Pilot scope", "Opportunity pipeline, account/contact layer, signal inbox, qualification scoring, RFQ intake, approval queue, weekly executive brief."),
        ("Deliverables", "Data model, working interface, agent workflows, pilot report, success metrics, and phase-two roadmap."),
        ("Ownership", "Level owns its data and generated operating memory. Code/licensing structure to be finalized in agreement."),
        ("Commercial options", "Fixed-fee pilot, build fee plus monthly support, strategy sprint then implementation, or long-term internal software partnership."),
    ]
    yy = y - 34
    for h, b in terms:
        c.setStrokeColor(colors.HexColor("#435260"))
        c.line(m + 10, yy + 5, pw - m - 10, yy + 5)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 6.7)
        c.drawString(m + 10, yy - 2, h)
        draw_wrapped_text(b, m + 118, yy - 2, cw - 130, size=6.5, leading=7.2, color=colors.HexColor("#dce3e8"), max_lines=1)
        yy -= 14

    c.setFillColor(MUTED)
    c.setFont("Helvetica", 7)
    c.drawString(m, 14, "Level Intelligence OS - one-page proposal sheet")
    c.drawRightString(pw - m, 14, "Draft for discussion")
    c.save()

def detailed_pdf():
    path = OUT / "Level Intelligence OS - Detailed Proposition.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.55 * inch,
    )

    story = []
    story.append(p("DETAILED PROPOSITION", "eyebrow"))
    story.append(p("Level Intelligence OS", "title"))
    story.append(p(
        "A private, agent-native operating layer that captures Level's analog construction, development, hospitality, and capital knowledge, structures messy language, routes agent-assisted decisions, and compounds from Level's own outcomes.",
        "lead",
    ))
    story.append(Rule(INK, 1.2))
    story.append(Spacer(1, 0.12 * inch))
    story.append(p("The Pain Before The Platform", "h1"))
    story.append(p(
        "The strongest case for Level Intelligence OS starts with the operating friction Level already feels: opportunity signals are scattered, relationship memory is trapped in people and inboxes, RFQs and permits are language-heavy, and executives spend too much attention assembling status instead of making decisions.",
        "lead",
    ))
    story.append(make_table([
        ["Pain Point", "What It Looks Like", "Why It Matters"],
        ["Scattered signals", "Permits, planning agendas, title changes, broker notes, franchise news, RFQs, and relationship updates arrive in different places.", "Promising opportunities can be seen late, duplicated, or missed entirely."],
        ["Trapped judgment", "Important context lives in individual memory, emails, spreadsheets, and project folders.", "Level's advantage is real but hard to search, hand off, measure, or improve."],
        ["RFQ and permit friction", "Packages, comments, deadlines, and missing requirements are buried in documents and emails.", "Teams spend time translating language into tasks and risk can surface late."],
        ["Status burden", "Leadership often needs to ask what changed, what is stuck, who owns it, and what needs approval.", "Executive time gets spent gathering context instead of exercising judgment."],
    ], [1.35 * inch, 2.55 * inch, 2.25 * inch]))
    story.append(PageBreak())

    story.append(p("Outcomes, Not Just Code", "h1"))
    story.append(p(
        "The deliverable is not a codebase sitting on a shelf. Code is the mechanism; operating outcomes are the product. Success should be measured by visible pipeline, faster triage, clearer approvals, fewer missed follow-ups, and enterprise memory that gets more valuable as Level uses it.",
        "lead",
    ))
    story.append(make_table([
        ["Outcome", "What Changes"],
        ["Pipeline visibility", "Active pursuits have owners, stages, scores, next actions, and aging risk in one operating view."],
        ["RFQ discipline", "Deadlines, missing requirements, scope risk, and go/no-go reasoning are extracted and routed."],
        ["Decision velocity", "Approvals, exceptions, and weekly executive briefs move leadership from status gathering to judgment."],
        ["Institutional memory", "Relationships, wins, losses, vendor outcomes, permit delays, and overrides stay inside the enterprise."],
        ["Adaptability", "Workflows can be added, omitted, or reshaped as Level changes because the operating layer is internal."],
    ], [1.75 * inch, 4.4 * inch]))
    story.append(Spacer(1, 0.18 * inch))
    story.append(make_table([
        ["Core Thesis", "What It Means"],
        ["Own the operating layer", "Level controls the domain model, data, workflows, and roadmap instead of forcing its work into static vendor software."],
        ["Start narrow", "The first build should focus on opportunity intelligence, RFQ workflow, approval queues, relationship memory, and executive briefs."],
        ["Compound over time", "Each signal, decision, win, loss, permit delay, vendor outcome, and relationship touch improves the next recommendation."],
    ], [1.6 * inch, 4.55 * inch]))
    story.append(p("1. Feasibility Assessment", "h1"))
    story.append(p(
        "This is feasible if scoped as an intelligence and operating layer first. It becomes risky only if framed as a full replacement for existing construction, accounting, CRM, or field systems before workflow evidence is gathered.",
        "lead",
    ))
    story.append(make_table([
        ["Question", "Assessment"],
        ["Can this be built?", "Yes. The first version can sit above current tools and coordinate signals, RFQs, relationships, approvals, and reporting."],
        ["What is the best first wedge?", "Opportunity pipeline, relationship memory, RFQ triage, go/no-go memos, approval queue, and weekly executive brief."],
        ["What are the main risks?", "Data quality, source access, unclear workflow ownership, and adoption. These are solvable through a narrow pilot and human review."],
        ["What should not be promised first?", "A full Procore, CRM, or accounting replacement. Integrate first; replace selectively only after evidence."],
    ], [1.75 * inch, 4.4 * inch]))
    story.append(Spacer(1, 0.18 * inch))
    story.append(p("Recommended proof points to add to the client narrative", "h2"))
    story.extend(bullet_items([
        "A sample weekly executive brief.",
        "A sample RFQ go/no-go memo.",
        "A simple ROI calculator using Level's actual volumes.",
        "A discovery checklist that shows what is needed to launch.",
        "A risk and mitigation table showing human control over consequential decisions.",
    ]))
    story.append(PageBreak())

    story.append(p("2. Operating Architecture", "h1"))
    story.append(p(
        "The system should be built as connected layers. Each layer can evolve independently as Level's markets, teams, and workflows change.",
        "lead",
    ))
    layer_rows = [
        ["Layer", "Function", "Examples"],
        ["Sources", "Observe external and internal work", "Permits, planning, title, RFQs, email, CRM, docs, project systems"],
        ["Ingestion", "Turn raw material into records", "Extract entities, dates, owners, deadlines, scopes, risks, confidence"],
        ["Domain model", "Represent Level's business", "Accounts, contacts, opportunities, sites, RFQs, bids, permits, approvals"],
        ["Memory", "Preserve structured and unstructured history", "Win/loss, relationship notes, vendor outcomes, permit cycles, decision logs"],
        ["Agents", "Prepare and route work", "Scout, qualify, summarize, compare, draft, flag, and brief"],
        ["Workflow and approval", "Control actions and risk", "Autonomous, suggest, approval required, manual"],
        ["Interfaces", "Make the system usable", "Executive home, signal inbox, pipeline, RFQ center, approval queue"],
    ]
    story.append(make_table(layer_rows, [1.25 * inch, 2.15 * inch, 2.75 * inch]))
    story.append(PageBreak())

    story.append(p("3. Opportunity Journey", "h1"))
    story.append(p(
        "The client should see a practical operating journey: market movement enters the system, agents prepare the context, humans make controlled decisions, and outcomes improve the next cycle.",
        "lead",
    ))
    journey = [
        ["Step", "What Happens", "Value"],
        ["1. Signal", "Permit, planning agenda, title transfer, broker note, brand expansion, or inbound RFQ is detected.", "Level sees movement earlier."],
        ["2. Enrich", "System connects site, owner, market, relationship history, jurisdiction, and similar work.", "Context is not trapped in inboxes or memory."],
        ["3. Qualify", "Opportunity is scored and explained using Level-specific criteria.", "Go/no-go decisions become more consistent."],
        ["4. Route", "The right owner receives the brief and next action.", "Work moves without status chasing."],
        ["5. Approve", "External outreach, bid submissions, vendor awards, budgets, and investor notes require approval.", "Agents help without creating risk."],
        ["6. Learn", "Wins, losses, delays, margin, response time, and overrides are retained.", "The system compounds from Level's own history."],
    ]
    story.append(make_table(journey, [0.9 * inch, 3.0 * inch, 2.25 * inch]))
    story.append(PageBreak())

    story.append(p("4. Agentic Structure and Governance", "h1"))
    story.append(p(
        "The system should use specialist agents with bounded tools, permissions, audit logs, and clear escalation rules. The design goal is useful autonomy without uncontrolled business risk.",
        "lead",
    ))
    agents = [
        ["Agent", "Role"],
        ["Orchestrator", "Routes requests, retrieves context, calls specialists, enforces permissions, and logs activity."],
        ["Signal Scout", "Detects signals, dedupes records, summarizes evidence, and proposes owner and next action."],
        ["Qualification", "Scores opportunities, explains reasoning, retrieves comparable history, and recommends pursue/watch/pass."],
        ["Relationship", "Builds account briefs, tracks owners, finds warm intros, flags stale relationships, and drafts outreach."],
        ["RFQ/Bid", "Summarizes packages, extracts requirements, detects missing documents, and drafts go/no-go memos."],
        ["Permit/Entitlement", "Tracks statuses, summarizes comments, drafts follow-ups, and flags jurisdictional blockers."],
        ["Executive Briefing", "Produces weekly briefs, pending decision lists, stuck pursuits, risk summaries, and market movement."],
    ]
    story.append(make_table(agents, [1.55 * inch, 4.6 * inch]))
    story.append(Spacer(1, 0.16 * inch))
    story.append(make_table([
        ["Authority Level", "Agent Can Do", "Examples"],
        ["Autonomous", "Act and log", "Summaries, classification, dedupe, internal briefs, score drafts"],
        ["Suggest", "Propose action", "Tasks, stage updates, draft emails, next actions, internal reminders"],
        ["Approval Required", "Prepare but wait for human approval", "External outreach, bid submission, vendor award, budget change, investor message"],
        ["Manual", "Assist only", "Contracts, financing, legal commitments, final strategic decisions"],
    ], [1.45 * inch, 2.1 * inch, 2.6 * inch]))
    story.append(PageBreak())

    story.append(p("5. ROI and Value Model", "h1"))
    story.append(p(
        "The initial ROI comes from time savings and fewer missed items. The stronger argument is better pursuit, stronger memory, faster decisions, and a proprietary operating dataset.",
        "lead",
    ))
    story.append(make_table([
        ["Value Lever", "How Value Is Created", "How To Measure"],
        ["Opportunity capture", "Signals become structured opportunities with owner, score, source, and next action.", "Signals captured, conversion rate, time to first action, wins by source"],
        ["Pursuit discipline", "Low-fit pursuits are filtered earlier while high-fit work gets faster attention.", "Go/no-go cycle time, win rate by score, avoided pursuit hours, margin by source"],
        ["Admin reduction", "Agents summarize, extract, draft, route, and remind across language-heavy workflows.", "Hours saved per RFQ, report time saved, overdue task reduction"],
        ["Risk reduction", "Permits, missing docs, stalled decisions, and deadline risk surface sooner.", "Permits past expected review, response time, missed deadline rate"],
        ["Enterprise asset", "Relationships, decisions, outcomes, vendors, and market patterns become private memory.", "Record completeness, recommendation accuracy, override rate, reuse of history"],
    ], [1.45 * inch, 2.95 * inch, 1.75 * inch]))
    story.append(PageBreak())

    story.append(p("6. 90-Day Pilot and Term Structure", "h1"))
    story.append(p(
        "The pilot should prove usefulness quickly while preserving a path to deeper workflows, integrations, and software ownership.",
        "lead",
    ))
    story.append(make_table([
        ["Period", "Focus", "Deliverable"],
        ["Weeks 1-2", "Foundation", "Data model, imports, roles, active pipeline, account/contact layer, opportunity views"],
        ["Weeks 3-5", "Intelligence", "Signal inbox, AI summaries, qualification score, relationship briefs, weekly digest"],
        ["Weeks 6-8", "Workflow", "RFQ intake, go/no-go memo, bid deadlines, approval queue, permit tracker foundation"],
        ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, stuck-decision reporting, outcome tracking, integration roadmap"],
    ], [1.25 * inch, 1.65 * inch, 3.25 * inch]))
    story.append(Spacer(1, 0.16 * inch))
    story.append(make_table([
        ["Term Area", "Recommended Position"],
        ["Engagement", "90-day pilot to design and build Level Intelligence OS v1."],
        ["Deliverables", "Data model, working interface, agent workflows, pilot report, success metrics, and phase-two roadmap."],
        ["Ownership", "Level owns its data and generated operating memory. Code/licensing structure to be finalized in agreement."],
        ["Commercial options", "Fixed-fee pilot, build fee plus monthly support, strategy sprint then implementation, or long-term internal software partnership."],
    ], [1.55 * inch, 4.6 * inch]))
    story.append(PageBreak())

    story.append(p("7. Why This Is Convincing", "h1"))
    story.append(p(
        "The strongest argument is not that AI can automate tasks. It is that Level can own a living software layer that gets more valuable as the company uses it.",
        "lead",
    ))
    story.append(make_table([
        ["Proof Wedge", "Trust Model", "Strategic Asset"],
        ["Pipeline clarity, RFQ triage, relationship briefs, approval queues, and executive reporting.", "Agents prepare decisions and explain reasoning. Humans approve money, outreach, vendors, legal, and investors.", "The system learns from Level's own opportunities, decisions, outcomes, and relationships."],
    ], [2.05 * inch, 2.05 * inch, 2.05 * inch]))
    story.append(Spacer(1, 0.18 * inch))
    story.append(p(
        "Recommended next step: run a focused discovery sprint to collect current pipeline, RFQ examples, relationship records, software map, and decision workflow. Use that to finalize the 90-day pilot scope and pricing.",
        "body",
    ))

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def docx_common_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor(58, 69, 81)
    for name, size in [("Heading 1", 22), ("Heading 2", 15), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(17, 22, 29)


def add_docx_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, val in enumerate(rows[0]):
        hdr[i].text = str(val)
        set_cell_shading(hdr[i], "17202B")
        set_cell_text_color(hdr[i], "FFFFFF")
        for pgh in hdr[i].paragraphs:
            for run in pgh.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def detailed_docx():
    doc = Document()
    docx_common_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Level Intelligence OS\n")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(17, 22, 29)
    sub = title.add_run("Detailed Proposition")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(13, 128, 108)

    doc.add_paragraph(
        "A private, agent-native operating layer that captures Level's analog construction, development, hospitality, and capital knowledge, structures messy language, routes agent-assisted decisions, and compounds from Level's own outcomes."
    )
    doc.add_heading("The Pain Before The Platform", level=1)
    doc.add_paragraph(
        "The case starts with Level's operating friction: opportunity signals are scattered, relationship memory is trapped in people and inboxes, RFQs and permits are language-heavy, and executives spend too much attention assembling status instead of making decisions."
    )
    add_docx_table(doc, [
        ["Pain Point", "What It Looks Like", "Why It Matters"],
        ["Scattered signals", "Permits, planning agendas, title changes, broker notes, franchise news, RFQs, and relationship updates arrive in different places.", "Promising opportunities can be seen late, duplicated, or missed entirely."],
        ["Trapped judgment", "Important context lives in individual memory, emails, spreadsheets, and project folders.", "Level's advantage is real but hard to search, hand off, measure, or improve."],
        ["RFQ and permit friction", "Packages, comments, deadlines, and missing requirements are buried in documents and emails.", "Teams spend time translating language into tasks and risk can surface late."],
        ["Status burden", "Leadership often needs to ask what changed, what is stuck, who owns it, and what needs approval.", "Executive time gets spent gathering context instead of exercising judgment."],
    ])
    doc.add_heading("Outcomes, Not Just Code", level=1)
    doc.add_paragraph(
        "The deliverable is not a codebase sitting on a shelf. Code is the mechanism; operating outcomes are the product. Success should be measured by visible pipeline, faster triage, clearer approvals, fewer missed follow-ups, and enterprise memory that gets more valuable as Level uses it."
    )
    add_docx_table(doc, [
        ["Outcome", "What Changes"],
        ["Pipeline visibility", "Active pursuits have owners, stages, scores, next actions, and aging risk in one operating view."],
        ["RFQ discipline", "Deadlines, missing requirements, scope risk, and go/no-go reasoning are extracted and routed."],
        ["Decision velocity", "Approvals, exceptions, and weekly executive briefs move leadership from status gathering to judgment."],
        ["Institutional memory", "Relationships, wins, losses, vendor outcomes, permit delays, and overrides stay inside the enterprise."],
        ["Adaptability", "Workflows can be added, omitted, or reshaped as Level changes because the operating layer is internal."],
    ])

    doc.add_heading("1. Feasibility Assessment", level=1)
    doc.add_paragraph(
        "This is feasible if scoped as an intelligence and operating layer first. It becomes risky only if framed as a full replacement for existing construction, accounting, CRM, or field systems before workflow evidence is gathered."
    )
    add_docx_table(doc, [
        ["Question", "Assessment"],
        ["Can this be built?", "Yes. The first version can sit above current tools and coordinate signals, RFQs, relationships, approvals, and reporting."],
        ["Best first wedge", "Opportunity pipeline, relationship memory, RFQ triage, go/no-go memos, approval queue, and weekly executive brief."],
        ["Main risks", "Data quality, source access, unclear workflow ownership, and adoption. These are solvable through a narrow pilot and human review."],
    ])

    doc.add_heading("2. Operating Architecture", level=1)
    add_docx_table(doc, [
        ["Layer", "Function", "Examples"],
        ["Sources", "Observe external and internal work", "Permits, planning, title, RFQs, email, CRM, docs, project systems"],
        ["Ingestion", "Turn raw material into records", "Extract entities, dates, owners, deadlines, scopes, risks, confidence"],
        ["Domain model", "Represent Level's business", "Accounts, contacts, opportunities, sites, RFQs, bids, permits, approvals"],
        ["Memory", "Preserve history", "Win/loss, relationship notes, vendor outcomes, permit cycles, decision logs"],
        ["Agents", "Prepare and route work", "Scout, qualify, summarize, compare, draft, flag, and brief"],
        ["Workflow", "Control actions and risk", "Autonomous, suggest, approval required, manual"],
        ["Interfaces", "Make the system usable", "Executive home, signal inbox, pipeline, RFQ center, approval queue"],
    ])

    doc.add_heading("3. Opportunity Journey", level=1)
    add_docx_table(doc, [
        ["Step", "What Happens", "Value"],
        ["Signal", "Permit, planning agenda, title transfer, broker note, brand expansion, or inbound RFQ is detected.", "Level sees movement earlier."],
        ["Enrich", "System connects site, owner, market, relationship history, jurisdiction, and similar work.", "Context is not trapped in inboxes or memory."],
        ["Qualify", "Opportunity is scored and explained using Level-specific criteria.", "Go/no-go decisions become more consistent."],
        ["Route", "The right owner receives the brief and next action.", "Work moves without status chasing."],
        ["Approve", "External outreach, bid submissions, vendor awards, budgets, and investor notes require approval.", "Agents help without creating risk."],
        ["Learn", "Wins, losses, delays, margin, response time, and overrides are retained.", "The system compounds from Level's own history."],
    ])

    doc.add_heading("4. Agentic Structure and Governance", level=1)
    add_docx_table(doc, [
        ["Agent", "Role"],
        ["Orchestrator", "Routes requests, retrieves context, calls specialists, enforces permissions, and logs activity."],
        ["Signal Scout", "Detects signals, dedupes records, summarizes evidence, and proposes owner and next action."],
        ["Qualification", "Scores opportunities, explains reasoning, retrieves comparable history, and recommends pursue/watch/pass."],
        ["Relationship", "Builds account briefs, tracks owners, finds warm intros, flags stale relationships, and drafts outreach."],
        ["RFQ/Bid", "Summarizes packages, extracts requirements, detects missing documents, and drafts go/no-go memos."],
        ["Permit/Entitlement", "Tracks statuses, summarizes comments, drafts follow-ups, and flags jurisdictional blockers."],
        ["Executive Briefing", "Produces weekly briefs, pending decision lists, stuck pursuits, risk summaries, and market movement."],
    ])
    add_docx_table(doc, [
        ["Authority Level", "Agent Can Do", "Examples"],
        ["Autonomous", "Act and log", "Summaries, classification, dedupe, internal briefs, score drafts"],
        ["Suggest", "Propose action", "Tasks, stage updates, draft emails, next actions, internal reminders"],
        ["Approval Required", "Prepare but wait for human approval", "External outreach, bid submission, vendor award, budget change, investor message"],
        ["Manual", "Assist only", "Contracts, financing, legal commitments, final strategic decisions"],
    ])

    doc.add_heading("5. ROI and Value Model", level=1)
    add_docx_table(doc, [
        ["Value Lever", "How Value Is Created", "How To Measure"],
        ["Opportunity capture", "Signals become structured opportunities with owner, score, source, and next action.", "Signals captured, conversion rate, time to first action, wins by source"],
        ["Pursuit discipline", "Low-fit pursuits are filtered earlier while high-fit work gets faster attention.", "Go/no-go cycle time, win rate by score, avoided pursuit hours, margin by source"],
        ["Admin reduction", "Agents summarize, extract, draft, route, and remind across language-heavy workflows.", "Hours saved per RFQ, report time saved, overdue task reduction"],
        ["Risk reduction", "Permits, missing docs, stalled decisions, and deadline risk surface sooner.", "Permits past expected review, response time, missed deadline rate"],
        ["Enterprise asset", "Relationships, decisions, outcomes, vendors, and market patterns become private memory.", "Record completeness, recommendation accuracy, override rate, reuse of history"],
    ])

    doc.add_heading("6. 90-Day Pilot and Term Structure", level=1)
    add_docx_table(doc, [
        ["Period", "Focus", "Deliverable"],
        ["Weeks 1-2", "Foundation", "Data model, imports, roles, active pipeline, account/contact layer, opportunity views"],
        ["Weeks 3-5", "Intelligence", "Signal inbox, AI summaries, qualification score, relationship briefs, weekly digest"],
        ["Weeks 6-8", "Workflow", "RFQ intake, go/no-go memo, bid deadlines, approval queue, permit tracker foundation"],
        ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, stuck-decision reporting, outcome tracking, integration roadmap"],
    ])

    doc.save(OUT / "Level Intelligence OS - Detailed Proposition.docx")


def one_page_docx():
    doc = Document()
    docx_common_styles(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)

    title = doc.add_paragraph()
    run = title.add_run("Level Intelligence OS")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(17, 22, 29)
    title.add_run("\nClient one-page dossier and term sheet").font.color.rgb = RGBColor(13, 128, 108)

    doc.add_paragraph(
        "A private, agent-native operating layer that converts Level X's analog construction, development, hospitality, and capital knowledge into a compounding digital asset."
    )
    add_docx_table(doc, [
        ["Pain Points"],
        ["Signals are scattered across permits, planning agendas, broker notes, franchise expansion, RFQs, and relationships. Judgment is trapped in people, inboxes, spreadsheets, and project memory. Leadership needs exceptions, decisions, risks, and next actions in one place."],
    ])
    add_docx_table(doc, [
        ["Outcomes, Not Just Code"],
        ["Code is the mechanism. The product is a working operating rhythm: visible pipeline, cleaner RFQ decisions, approval rails, weekly executive briefs, and compounding operating memory."],
    ])
    add_docx_table(doc, [
        ["Feasibility", "Risk", "Best Wedge"],
        ["High if scoped as an intelligence and workflow layer across current tools.", "Data quality and adoption. Solve with a narrow pilot and human review.", "Opportunity pipeline, relationship memory, RFQ triage, and weekly executive brief."],
    ])
    add_docx_table(doc, [
        ["What Level Gets", "Why It Matters"],
        ["Earlier signals, Level-specific scoring, agent workflows, and human approval rails.", "Internal software can evolve through language and compound from every decision and outcome."],
    ])
    add_docx_table(doc, [
        ["Value Prop", "Meaning"],
        ["Faster pursuit", "Signals become opportunities with owners, scores, and next actions."],
        ["Better margin", "Go/no-go discipline reduces wasted pursuits and highlights risk earlier."],
        ["Lower admin load", "Agents summarize, extract, draft, route, and remind across language-heavy work."],
        ["Enterprise asset", "Relationships, decisions, and outcomes become proprietary operating memory."],
    ])
    add_docx_table(doc, [
        ["Term Area", "Skeleton"],
        ["Engagement", "90-day pilot to design and build Level Intelligence OS v1."],
        ["Pilot scope", "Opportunity pipeline, account/contact layer, signal inbox, qualification scoring, RFQ intake, approval queue, weekly executive brief."],
        ["Ownership", "Level owns its data and generated operating memory. Code/licensing structure to be finalized in agreement."],
        ["Commercial options", "Fixed-fee pilot, build fee plus monthly support, strategy sprint then implementation, or long-term internal software partnership."],
    ])
    doc.save(OUT / "Level Intelligence OS - One Page Dossier and Term Sheet.docx")


if __name__ == "__main__":
    one_page_pdf()
    detailed_pdf()
    detailed_docx()
    one_page_docx()
    print(f"Generated polished exports in {OUT}")
