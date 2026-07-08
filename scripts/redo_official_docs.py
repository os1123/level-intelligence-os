from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/omar/Grounded/Projects/Level Intelligence OS")
OUT = ROOT / "official"

INK = "121820"
DARK = "17202B"
MUTED = "5B6775"
LINE = "D8DDE3"
PAPER = "F7F7F4"
ACCENT = "0E806C"
YELLOW = "C99B2E"
WHITE = "FFFFFF"
LIGHT_GREEN = "E7F3EF"
LIGHT_YELLOW = "F6EEDA"


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=LINE, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct=100):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct * 50))
    tbl_w.set(qn("w:type"), "pct")


def style_doc(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(MUTED)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in [
        ("Title", 30, INK),
        ("Heading 1", 18, INK),
        ("Heading 2", 12, INK),
        ("Heading 3", 10.5, INK),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = name != "Title"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(8 if name == "Heading 1" else 4)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = False
        st.paragraph_format.keep_together = False
        st.paragraph_format.page_break_before = False


def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Aptos"
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


def cell_text(cell, text, size=8.5, color=MUTED, bold=False, align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins(cell)
    set_cell_border(cell)
    return p


def add_kicker(paragraph, text):
    run = paragraph.add_run(text.upper())
    set_run(run, size=7.5, color=ACCENT, bold=True)


def add_title_block(doc, title, subtitle, kicker="official client document"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_kicker(p, kicker)
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(title)
    set_run(r, size=28, color=INK, bold=False)
    s = doc.add_paragraph()
    s.paragraph_format.space_after = Pt(8)
    sr = s.add_run(subtitle)
    set_run(sr, size=10.5, color=MUTED)


def add_callout(doc, title, body, fill=DARK, title_color=WHITE, body_color="E2E8EE"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_shading(cell, fill)
    set_cell_border(cell, fill, "4")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, size=11, color=title_color, bold=True)
    b = cell.add_paragraph()
    b.paragraph_format.space_after = Pt(0)
    br = b.add_run(body)
    set_run(br, size=9, color=body_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows, widths=None, font_size=8.4, header=True):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            is_header = header and r_idx == 0
            if is_header:
                set_shading(cells[c_idx], DARK)
                cell_text(cells[c_idx], str(value), size=8, color=WHITE, bold=True)
            else:
                set_shading(cells[c_idx], WHITE if r_idx % 2 else PAPER)
                cell_text(cells[c_idx], str(value), size=font_size, color=MUTED)
            if widths and c_idx < len(widths):
                cells[c_idx].width = Inches(widths[c_idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def bullet_list(doc, items, size=9):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(item)
        set_run(r, size=size, color=MUTED)


def add_section(doc, title, lead=None):
    p = doc.add_heading(title, level=1)
    p.paragraph_format.keep_with_next = False
    if lead:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        r = para.add_run(lead)
        set_run(r, size=10, color=MUTED)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    r = hp.add_run("Level Intelligence OS")
    set_run(r, size=8.5, color=INK, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    left = fp.add_run("Draft for discussion")
    set_run(left, size=8, color=MUTED)
    fp.add_run("\t")
    set_page_number(fp)


def build_one_pager():
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.28)
    section.left_margin = Inches(0.38)
    section.right_margin = Inches(0.38)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_kicker(p, "official one-page dossier")
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Level Intelligence OS")
    set_run(run, size=25, color=INK)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(5)
    sr = sub.add_run("A private construction intelligence operating system built around Level's relationships, RFQs, approvals, Procore reality, and operating judgment.")
    set_run(sr, size=9.5, color=MUTED)

    top = doc.add_table(rows=1, cols=3)
    set_table_width(top)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    top.autofit = True
    pain = [
        ("Scattered signals", "Permits, RFQs, planning activity, broker notes, franchise news, emails, and relationship tips arrive in different places."),
        ("Trapped judgment", "Fit, risk, margin, timing, relationships, and prior lessons live in people, inboxes, calls, and informal notes."),
        ("Status burden", "Leadership spends time reconstructing what changed, what is stuck, who owns it, and what needs approval."),
    ]
    for i, (h, b) in enumerate(pain):
        cell = top.cell(0, i)
        set_shading(cell, PAPER)
        set_cell_border(cell)
        set_cell_margins(cell, top=120, bottom=110, start=140, end=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        hr = p.add_run(h)
        set_run(hr, size=9.5, color=INK, bold=True)
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(0)
        br = bp.add_run(b)
        set_run(br, size=7.4, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    mid = doc.add_table(rows=1, cols=2)
    set_table_width(mid)
    mid.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = mid.cell(0, 0)
    right = mid.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=125, bottom=120, start=150, end=150)
        set_cell_border(cell)
    set_shading(left, DARK)
    set_shading(right, WHITE)
    p = left.paragraphs[0]
    r = p.add_run("We give you outcomes, not just code.")
    set_run(r, size=12.5, color=WHITE, bold=True)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Code is the mechanism. The product is cleaner pursuit visibility, faster RFQ triage, fewer missed follow-ups, a weekly executive decision rhythm, and Level-owned operating memory that compounds.")
    set_run(r2, size=8.2, color="E2E8EE")

    p = right.paragraphs[0]
    r = p.add_run("Built for Level, wrapped around Procore first.")
    set_run(r, size=12.5, color=INK, bold=True)
    p2 = right.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("This is not a vendor clone or a rip-and-replace pitch. Procore remains the current execution source of truth while Level Intelligence OS becomes the intelligence, RFQ, relationship, approval, and executive layer above it.")
    set_run(r2, size=8.2, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    cols = doc.add_table(rows=1, cols=3)
    set_table_width(cols)
    blocks = [
        ("What Level gets", [
            "Signal inbox for early project, permit, planning, ownership, broker, and RFQ movement.",
            "Level-specific qualification score: pursue, watch, or pass with reasons.",
            "RFQ command center with deadlines, missing items, risks, and go/no-go memo.",
            "Executive brief showing what changed, what is stuck, and what needs judgment.",
        ], LIGHT_GREEN),
        ("90-day pilot", [
            "Weeks 1-2: data model, roles, active pursuits, account/contact layer.",
            "Weeks 3-5: signal inbox, opportunity summaries, scoring, relationship briefs.",
            "Weeks 6-8: RFQ command center, deadline extraction, approval queue.",
            "Weeks 9-12: weekly brief, outcome capture, adoption review, phase-two roadmap.",
        ], WHITE),
        ("Term sheet skeleton", [
            "Engagement: 90-day private Level Intelligence OS pilot.",
            "Ownership: Level owns data, operating memory, domain model, workflow definitions, and exportable records.",
            "Commercial: fixed-fee pilot, build plus retainer, or long-term internal software partnership.",
            "Control: external outreach, bid submission, vendor award, budget, investor, legal, and financing moves require human approval.",
        ], LIGHT_YELLOW),
    ]
    for i, (h, items, fill) in enumerate(blocks):
        cell = cols.cell(0, i)
        set_shading(cell, fill)
        set_cell_border(cell)
        set_cell_margins(cell, top=120, bottom=110, start=130, end=130)
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(h)
        set_run(hr, size=10.5, color=INK, bold=True)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.08)
            p.paragraph_format.first_line_indent = Inches(-0.08)
            p.paragraph_format.space_after = Pt(1)
            rr = p.add_run("- " + item)
            set_run(rr, size=7.25, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    value = doc.add_table(rows=1, cols=4)
    set_table_width(value)
    value_blocks = [
        ("Opportunity capture", "Earlier signals become owned opportunities with source, score, owner, and next action."),
        ("Decision speed", "RFQ facts, blockers, and approval needs surface before they become bottlenecks."),
        ("Pursuit discipline", "Level spends less time on poor-fit work and more time on high-fit relationships."),
        ("Compounding memory", "Wins, losses, overrides, and outcomes improve the next recommendation."),
    ]
    for i, (h, b) in enumerate(value_blocks):
        cell = value.cell(0, i)
        set_shading(cell, WHITE)
        set_cell_border(cell)
        set_cell_margins(cell, top=95, bottom=90, start=110, end=110)
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run(h)
        set_run(hr, size=8.8, color=INK, bold=True)
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(0)
        br = bp.add_run(b)
        set_run(br, size=6.9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Close: Level should not rent generic intelligence when it can build the operating memory that makes its own team sharper.")
    set_run(r, size=8.2, color=INK, bold=True)

    path = OUT / "Level Intelligence OS - Official One Pager.docx"
    doc.save(path)
    return path


def build_comprehensive_docx():
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.64)
    section.right_margin = Inches(0.64)
    add_header_footer(section)

    add_title_block(
        doc,
        "Level Intelligence OS",
        "Comprehensive client packet for a private construction intelligence operating system built around Level's workflows, Procore reality, and operating judgment.",
        "comprehensive client packet",
    )
    add_callout(
        doc,
        "Commercial promise: outcomes, not just code",
        "Code is the mechanism. The product is operating lift: cleaner opportunity visibility, faster RFQ triage, stronger go/no-go discipline, fewer missed follow-ups, a weekly executive decision rhythm, and Level-owned operating memory that compounds over time.",
    )

    add_table(doc, [
        ["Packet Section", "Purpose"],
        ["1. Executive thesis", "Why Level should own a private operating layer."],
        ["2. Pain before platform", "Where scattered judgment creates drag."],
        ["3. Built for Level", "Why this is not a vendor clone or replacement pitch."],
        ["4. Proven infrastructure", "Reusable construction-platform primitives that lower build risk."],
        ["5. Procore bridge", "How to wrap around the current source of truth before replacing anything."],
        ["6. Market shift", "Why software is moving toward forward-deployed, company-specific systems."],
        ["7. Operating design", "How find, qualify, route, convert, and learn becomes software."],
        ["8. Pilot and terms", "How to prove value in 90 days and structure the engagement."],
    ], [1.7, 4.7], font_size=8.5)

    doc.add_page_break()

    add_section(doc, "1. Executive Thesis", "Level already has the hard part: trusted relationships, construction judgment, development instincts, hospitality experience, capital discipline, vendor memory, and practical pattern recognition about which opportunities are worth time.")
    para = doc.add_paragraph("The opportunity is to convert that judgment into a private operating system Level owns. This is not a generic AI layer, chatbot, dashboard, CRM migration, or forced Procore replacement. It is an internal operating layer around business development, preconstruction, project pursuit, development, hospitality, capital workflows, and executive management.")
    set_run(para.runs[0], size=9.5, color=MUTED)
    add_table(doc, [
        ["Near-Term Outcome", "Long-Term Strategic Asset"],
        ["Cleaner active pipeline, faster RFQ review, better go/no-go discipline, fewer missed follow-ups, and weekly executive visibility.", "A Level-owned operating memory that improves pursuit, relationship leverage, execution handoff, and future software adaptability."],
    ], [3.1, 3.1], font_size=8.6)
    add_callout(doc, "Core proposition", "Level should not rent generic intelligence when it can build the operating memory that makes its own team sharper.", fill=LIGHT_GREEN, title_color=INK, body_color=MUTED)

    add_section(doc, "2. The Pain Before The Platform", "The highest-value information in construction pursuit work rarely lives in one system. It lives in RFQs, emails, calls, spreadsheets, permit threads, broker notes, project folders, meeting notes, personal memory, and executive judgment.")
    add_table(doc, [
        ["Pain Point", "What It Looks Like", "Why It Matters"],
        ["Scattered signals", "Permits, planning activity, ownership changes, broker notes, franchise news, RFQs, and relationship tips arrive in different places.", "Promising opportunities can be seen late, duplicated, or missed."],
        ["Trapped judgment", "Fit, risk, timing, margin, and relationship context live in individual memory, inboxes, and informal notes.", "Level's advantage is real but hard to search, compare, hand off, or improve."],
        ["RFQ friction", "Deadlines, missing documents, addenda, site risk, scope ambiguity, and submission requirements are buried in language-heavy packages.", "Teams spend time translating documents into tasks and may surface risk late."],
        ["Status burden", "Leadership needs to ask what changed, what is stuck, who owns it, and what needs approval.", "Executive time is spent assembling context instead of exercising judgment."],
    ], [1.3, 2.7, 2.2], font_size=8.1)

    add_section(doc, "3. Built For Level, Not For The Category", "Mercator and similar tools are useful category references because they show that earlier construction intelligence has value. But Level Intelligence OS is not a Mercator replacement pitch and not a vendor comparison exercise.")
    add_table(doc, [
        ["Generic Vendor Logic", "Level Intelligence OS Logic"],
        ["Expose market or project signals.", "Expose signals and translate them into Level-specific opportunities, owners, scores, and next actions."],
        ["Use vendor-defined fields and workflows.", "Represent Level's own domain model: accounts, contacts, sites, RFQs, permits, bids, approvals, relationships, and decisions."],
        ["Optimize for broad category adoption.", "Optimize for Level's operating judgment, pursuit discipline, leadership rhythm, and long-term ownership."],
        ["Data remains mainly a tool output.", "Data becomes Level-owned operating memory that compounds over time."],
    ], [2.7, 3.6], font_size=8.3)
    add_callout(doc, "Positioning sentence", "Level Intelligence OS is not built to compete with anyone else's product. It is built to make Level's own operating model software-native.", fill=LIGHT_YELLOW, title_color=INK, body_color=MUTED)

    add_section(doc, "4. Proven Infrastructure Patterns We Bring", "We are not starting from an empty page. The build can reuse internal construction-platform patterns already worked through in prior platform development while keeping the client-facing system fully Level-specific.")
    add_table(doc, [
        ["Pattern", "How It Applies To Level"],
        ["Construction domain model", "Accounts, contacts, sites, projects, permits, RFQs, bids, budget items, approvals, outcomes, and operating notes become connected records."],
        ["Workflow state machines", "Opportunities, RFQs, permits, approvals, and later vendor/payment workflows move through explicit statuses instead of scattered messages."],
        ["Specialist agent structure", "Orchestrated agents handle bounded jobs such as signal intake, qualification, RFQ extraction, permit monitoring, executive briefing, and finance context."],
        ["Human-in-the-loop controls", "Authority levels, approval queues, audit logs, escalation rules, and explicit human review keep consequential actions under Level control."],
        ["Document intelligence", "RFQs, permits, invoices, emails, addenda, and project narratives become deadlines, risks, tasks, missing information, and decision memos."],
        ["Feedback loops", "Wins, losses, overrides, delays, vendor outcomes, and cycle times become signals for better future recommendations."],
    ], [1.65, 4.65], font_size=8.2)

    add_section(doc, "5. Procore Bridge Strategy", "Level Intelligence OS should wrap around Procore first, not replace it. Procore remains the execution source of truth while Level Intelligence OS becomes the intelligence, pursuit, RFQ, relationship, approval, and executive layer around it.")
    add_table(doc, [
        ["Phase", "Position", "Outcome"],
        ["1. Read and reference", "Treat Procore as the current project source of truth. Pull or reference only the data needed for visibility and decisions.", "No disruption to field or project teams."],
        ["2. Workflow overlay", "Use Level Intelligence OS for signals, RFQs, go/no-go memos, relationship memory, approvals, and executive briefs.", "Level gains operating rhythm without forcing a platform migration."],
        ["3. Evidence-based rationalization", "Measure which workflows are better handled inside Level's own system and which should stay in Procore.", "Software decisions become based on observed value."],
        ["4. Selective replacement", "Move modules away from monolithic software only when Level's owned layer is clearly better, safer, and cheaper to operate.", "Dependency reduces over time without risky rip-and-replace."],
    ], [1.25, 2.7, 2.2], font_size=8.0)
    add_callout(doc, "Operating principle: earn the right to replace", "The long-term goal is less reliance on monolithic software, but the first move is a Procore-first intelligence layer that creates measurable value now.", fill=DARK, title_color=WHITE, body_color="E2E8EE")

    doc.add_page_break()

    add_section(doc, "6. Market Shift: Forward-Deployed Software", "The software market is moving from fixed SaaS screens toward company-specific AI systems deployed inside real workflows. OpenAI's 2026 Deployment Company and forward-deployed engineering model are a clear market signal: enterprise value increasingly comes from fitting intelligence into real operating constraints, not simply licensing a generic tool.")
    add_table(doc, [
        ["Shift", "Old Pattern", "New Pattern"],
        ["Static SaaS to internal systems", "The company adapts to vendor-defined screens and workflows.", "Software adapts to the company's language, judgment, data, approvals, and operating rhythm."],
        ["Generic automation to agentic work", "Rules handle narrow tasks but struggle with messy language-heavy work.", "Agents read, synthesize, route, draft, and learn from RFQs, emails, permits, notes, and decisions."],
        ["Vendor lock-in to owned memory", "Operating history lives inside external tools and exports.", "The company's decisions, outcomes, and workflows become a proprietary asset."],
        ["Implementation to forward deployment", "Software is installed and handed off.", "Experts map work, ship inside real constraints, tune with operators, and turn adoption into measurable outcomes."],
        ["Manual change requests to language-native change", "New workflows require slow product cycles or vendor roadmaps.", "Teams can describe additions, omissions, and changes in natural language, then convert them into owned software modules."],
    ], [1.45, 2.25, 2.55], font_size=8.0)
    note = doc.add_paragraph("Source note: OpenAI launched the OpenAI Deployment Company on May 11, 2026 and describes forward deployed engineering as a way to bring AI into production for complex real-world use cases.")
    set_run(note.runs[0], size=8, color=MUTED, italic=True)

    add_section(doc, "7. What The System Does", "Level Intelligence OS converts messy pursuit work into an operating loop: find, qualify, route, convert, and learn.")
    add_table(doc, [
        ["Capability", "What It Does", "Client-Facing Outcome"],
        ["Find", "Detects public and private project signals before they become obvious.", "Level sees more opportunities earlier."],
        ["Qualify", "Scores opportunities using Level's criteria, history, relationships, and capacity.", "Leadership spends less time on poor-fit work."],
        ["Route", "Assigns owners, next actions, approvals, and follow-ups.", "Good opportunities do not stall in inboxes."],
        ["Convert", "Supports RFQs, go/no-go memos, pursuit briefs, and executive reporting.", "Level turns intelligence into action."],
        ["Learn", "Captures decisions, overrides, wins, losses, delays, and outcomes.", "Every cycle improves the next score, brief, and recommendation."],
    ], [1.0, 2.65, 2.55], font_size=8.2)

    doc.add_page_break()

    add_section(doc, "8. Product Modules")
    add_table(doc, [
        ["Module", "What It Includes", "Why It Matters"],
        ["Signal inbox", "Permits, planning agendas, title/ownership changes, rezoning, franchise expansion, broker notes, inbound RFQs, CRM/email activity.", "Early signals become reviewable records instead of scattered noise."],
        ["Opportunity graph", "Accounts, contacts, sites, jurisdictions, RFQs, bids, permits, decisions, owners, next actions.", "Level sees relationships between people, projects, sites, and pursuits."],
        ["Qualification engine", "Pursue/watch/pass, confidence, reasoning, missing info, risks, next action, owner, comparable history.", "Level's own standards become repeatable without replacing human judgment."],
        ["Relationship memory", "Relationship owner, last touch, prior work, warm intros, next follow-up, leadership brief.", "Relationship context becomes searchable and actionable."],
        ["RFQ command center", "Scope summary, deadlines, addenda, required attachments, missing documents, risks, go/no-go memo, bid tasks.", "RFQs become decision-ready work."],
        ["Approval queue", "Recommendations, evidence, confidence, approver, decision, outcome.", "Agents prepare work while humans control consequential actions."],
        ["Executive brief", "Top signals, active pursuits, stalled items, RFQs due, approvals needed, relationship follow-ups, lessons learned.", "Leadership gets a weekly decision rhythm from live system data."],
    ], [1.3, 3.0, 2.0], font_size=7.8)

    add_section(doc, "9. Agentic Operating Model", "The system should be composed of specialist workflows under one governance model. This avoids the failure mode of one vague assistant that is expected to know everything.")
    add_table(doc, [
        ["Workflow Agent", "Role"],
        ["Signal Scout", "Finds, dedupes, summarizes, and routes project signals."],
        ["Qualification", "Scores opportunities against Level-specific criteria and explains the recommendation."],
        ["Relationship", "Builds account/contact briefs, tracks owners, suggests follow-ups, and supports warm intros."],
        ["RFQ/Bid", "Extracts requirements, deadlines, risks, missing documents, addenda, and go/no-go memos."],
        ["Permit/Entitlement", "Tracks jurisdiction status, review comments, days in review, and hidden timing risk."],
        ["Executive Briefing", "Turns system activity into a weekly decision brief."],
        ["Finance/Capital", "Connects opportunities to financing, underwriting, investor context, and capital constraints where relevant."],
    ], [1.6, 4.7], font_size=8.2)

    doc.add_page_break()

    add_section(doc, "10. Human Control And Trust")
    add_table(doc, [
        ["Authority Level", "System Can Do", "Human Control"],
        ["Autonomous", "Classify signals, summarize RFQs, dedupe contacts, draft internal briefs.", "Logged for review."],
        ["Suggest", "Recommend owner, next action, pursue/watch/pass, follow-up draft.", "Human can accept, edit, or reject."],
        ["Approval Required", "External outreach, bid submission, vendor award, budget change, investor communication.", "Explicit approval required."],
        ["Manual", "Legal, contracts, financing, final strategic decisions.", "Human executes."],
    ], [1.35, 3.0, 1.85], font_size=8.2)

    add_section(doc, "11. ROI And Value Model", "The ROI should not be sold only as labor savings. The bigger value is better pursuit discipline, faster decisions, fewer blind spots, stronger relationship leverage, and proprietary operating memory.")
    add_table(doc, [
        ["Value Lever", "How Value Is Created", "How To Measure"],
        ["Opportunity capture", "Signals become structured opportunities with owner, score, source, and next action.", "Signals captured, time to first action, conversion by source, wins by source."],
        ["Pursuit discipline", "Low-fit pursuits are filtered earlier while high-fit work receives faster attention.", "Go/no-go cycle time, win rate by score, avoided pursuit hours, margin by source."],
        ["Admin reduction", "Agents summarize, extract, draft, route, and remind across language-heavy workflows.", "Hours saved per RFQ, weekly report time saved, overdue task reduction."],
        ["Risk reduction", "Permits, missing documents, stalled decisions, and deadline risk surface sooner.", "Missed deadline rate, response time, unresolved blocker age."],
        ["Enterprise asset", "Relationships, decisions, outcomes, vendors, and market patterns become private memory.", "Record completeness, reuse of history, recommendation accuracy, override rate."],
    ], [1.35, 3.0, 1.85], font_size=8.0)
    add_callout(doc, "Compounding logic", "Each signal, decision, win, loss, permit delay, vendor outcome, and relationship touch improves the next recommendation.", fill=LIGHT_GREEN, title_color=INK, body_color=MUTED)

    doc.add_page_break()

    add_section(doc, "12. 90-Day Pilot", "The pilot should prove usefulness in workflows Level already runs: opportunity review, RFQ triage, relationship follow-up, approval routing, and executive reporting.")
    add_table(doc, [
        ["Period", "Focus", "Client-Visible Deliverable"],
        ["Weeks 1-2", "Foundation", "Data model, roles, active opportunities, account/contact layer, executive dashboard skeleton."],
        ["Weeks 3-5", "Intelligence", "Signal inbox, opportunity summaries, qualification scoring, relationship briefs, first weekly digest."],
        ["Weeks 6-8", "Workflow", "RFQ command center, deadline extraction, go/no-go memo, approval queue, bid task tracking."],
        ["Weeks 9-12", "Operating rhythm", "Weekly executive brief, outcome tracking, adoption review, phase-two roadmap."],
    ], [1.15, 1.55, 3.5], font_size=8.3)
    doc.add_heading("Pilot Success Metrics", level=2)
    bullet_list(doc, [
        "100 percent of active pursuits visible in the system.",
        "Each pursuit has owner, stage, source, score, next action, and age.",
        "At least 10 useful signals captured, reviewed, and routed.",
        "RFQs tracked with deadlines, requirements, and go/no-go status.",
        "Weekly executive brief generated from live system data.",
        "Leadership identifies faster decisions, fewer blind spots, or clearer accountability.",
        "Phase-two roadmap is based on observed value, not assumptions.",
    ], size=8.7)

    add_section(doc, "13. Phase Two And System Expansion", "Expansion should be based on where Level sees the highest operating value during the pilot, not a preset software wish list.")
    add_table(doc, [
        ["Expansion Area", "Potential Scope"],
        ["Permit and entitlement monitoring", "Jurisdiction status, comments, timing risk, follow-up drafts, days in review."],
        ["Email, calendar, CRM, storage integration", "Source access, relationship context, meeting history, follow-up capture."],
        ["Procore read-only bridge", "Treat Procore as the current project source of truth while Level Intelligence OS adds intelligence, pursuit, and executive workflow above it."],
        ["Vendor and subcontractor memory", "Performance history, issue patterns, pricing, reliability, project outcomes."],
        ["Project handoff", "Transfer pursuit context into execution planning and kickoff."],
        ["Finance and capital workflow", "Underwriting context, investor notes, financing constraints, capital approval support."],
        ["Hospitality pipeline", "Brand/operator tracking, franchise expansion, site qualification, relationship history."],
        ["Software rationalization", "Usage audit to decide what should remain, integrate, downgrade, or eventually be replaced by Level-owned modules."],
    ], [2.0, 4.25], font_size=8.0)

    doc.add_page_break()

    add_section(doc, "14. Commercial, Ownership, And Trust Principles")
    add_table(doc, [
        ["Area", "Recommended Position"],
        ["Engagement", "90-day pilot to design, build, test, and operationalize Level Intelligence OS v1."],
        ["Pilot deliverables", "Workflow map, data model, opportunity/account/contact database, signal inbox, scoring framework, Procore-first bridge strategy, RFQ workflow, approval queue, weekly executive brief, success report, phase-two roadmap."],
        ["Ownership", "Level owns its data, operating memory, domain model, workflow definitions, and exportable generated records. Code ownership/licensing to be defined in the final agreement."],
        ["Model infrastructure", "AI model providers remain replaceable infrastructure, not the strategic asset."],
        ["Confidentiality", "Level data is confidential. No external reuse without written permission. Role-based access and audit logs are required."],
        ["Human approval", "External outreach, bid submission, vendor award, budget change, investor communication, legal, financing, and final strategic decisions require human control."],
        ["Commercial options", "Fixed-fee pilot, build fee plus monthly operating retainer, strategy/design sprint followed by implementation, or long-term internal software partnership."],
    ], [1.55, 4.7], font_size=8.0)

    add_section(doc, "15. Discovery Agenda", "The first discovery session should focus on actual workflow, not abstract AI strategy.")
    add_table(doc, [
        ["Discovery Topic", "Questions To Answer"],
        ["Opportunity sources", "Where do new opportunities originate today? Which sources matter most?"],
        ["Qualification criteria", "What makes a Level opportunity attractive or unattractive? Who decides?"],
        ["RFQ workflow", "Where do RFQs live? How are deadlines, addenda, tasks, and go/no-go decisions handled?"],
        ["Relationship memory", "Which accounts, contacts, owners, brands, developers, and vendors matter most?"],
        ["Executive reporting", "What does leadership need to see every week? What is currently hard to assemble?"],
        ["Systems and data", "Which tools are in use today: Procore, CRM, email, calendar, accounting, storage? What should be read, synced, imported, or left alone?"],
        ["Approval rules", "Which actions require explicit approval? Who approves them?"],
        ["Pilot success", "What would make the pilot obviously worth continuing?"],
    ], [1.7, 4.55], font_size=8.0)
    add_callout(doc, "Recommended next step", "Run a 60-minute workflow discovery session, collect sample RFQs and current active pursuits, then finalize the pilot scope and commercial structure.", fill=DARK, title_color=WHITE, body_color="E2E8EE")

    doc.add_page_break()

    add_section(doc, "16. Closing Position", "Level Intelligence OS should be presented as a serious operating asset, not an AI experiment.")
    para = doc.add_paragraph("The first version gives Level one place to see opportunities, understand fit, track RFQs, preserve relationship memory, and route decisions. The longer-term version becomes Level's private construction intelligence layer: a system that learns from Level's signals, relationships, decisions, wins, losses, and operating judgment.")
    set_run(para.runs[0], size=10, color=MUTED)
    add_callout(doc, "Final message", "Level can build internal software that grows with the company, adapts through language, and compounds from its own operating history.", fill=LIGHT_GREEN, title_color=INK, body_color=MUTED)
    add_table(doc, [
        ["Immediate Ask", "Approve a focused 90-day pilot."],
        ["What Level Gets", "A working operating rhythm, not a codebase sitting on a shelf."],
        ["Why It Matters", "The software becomes more valuable because it learns Level."],
    ], [1.65, 4.6], font_size=8.8, header=False)

    path = OUT / "Level Intelligence OS - Comprehensive Client Packet.docx"
    doc.save(path)
    return path


def main():
    OUT.mkdir(exist_ok=True)
    one = build_one_pager()
    packet = build_comprehensive_docx()
    print(one)
    print(packet)


if __name__ == "__main__":
    main()
